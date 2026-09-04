# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import doc_common as dc

doc = Document()
dc.set_base_style(doc)
dc.add_page_number_footer(doc)
dc.enable_auto_update_fields(doc)

dc.add_title_page(
    doc,
    "PX4 기반 틸트로터 VTOL 비행제어시스템",
    "HILS/시각화 연동 인터페이스 통제 문서 (ICD)",
    doc_no="ICD-PXTR-HILS-001",
    rev="A (초안)",
    date_str="2026-08-31",
)

dc.add_revision_history(doc, [
    ["A", "2026-08-31", "최초 발행 (초안) — 표준 HIL 메시지 및 커스텀 틸트 메시지 정의", "비행제어SW팀"],
])

doc.add_heading("목차 (Table of Contents)", level=1)
dc.add_toc(doc)
doc.add_page_break()

# =====================================================================
doc.add_heading("1. 개요", level=1)
doc.add_heading("1.1 목적", level=2)
doc.add_paragraph(
    "본 문서는 「EICD-PXTR-HILS-001」에서 정의한 외부 인터페이스(EICD-01, EICD-02)에 대해, 실제로 "
    "송수신되는 메시지의 프레임 구조, 필드, 자료형, 단위, 유효범위를 소프트웨어 구현 수준에서 정의한다. "
    "본 문서를 근거로 HILS 환경 시뮬레이터(ENV) 및 시각화 시스템(VIS) 개발자는 FCC(PX4 OFP)와 데이터를 "
    "교환하는 소프트웨어를 구현할 수 있다."
)

doc.add_heading("1.2 적용범위", level=2)
doc.add_paragraph(
    "본 문서는 아래 4개 논리 채널에 적용된다."
)
dc.add_table(doc, ["채널", "방향", "목적", "관련 EICD"], [
    ["A. 센서/피드백 주입", "ENV → FCC", "가상 IMU/GPS/기압/자기장 및 틸트 서보 피드백을 FCC에 주입", "EICD-01"],
    ["B. 액추에이터/틸트 명령", "FCC → ENV", "모터 추력 및 틸트 목표각 명령을 ENV의 동역학 모델에 전달", "EICD-01"],
    ["C. 상태 텔레메트리", "FCC → VIS", "위치/속도/자세/풍향/틸트상태 등 시각화용 상태값 송출", "EICD-02"],
    ["D. 조종입력 검증", "RC/GCS → FCC", "조종기 입력을 FCC에 전달하여 조종 반응 검증", "EICD-05"],
], col_widths_cm=[4.0, 2.5, 6.5, 3.0])

doc.add_heading("1.3 참고 문서", level=2)
for t in [
    "MAVLink Common Message Set — https://mavlink.io/en/messages/common.html",
    "MAVLink Development Message Set — https://mavlink.io/en/messages/development.html",
    "PX4 uORB 메시지 정의 (airbility-dev/px4-custom-firmware, msg/*.msg)",
    "src/modules/simulation/simulator_mavlink/SimulatorMavlink.cpp (HIL 메시지 처리 레퍼런스 구현)",
    "EICD-PXTR-HILS-001 외부 인터페이스 통제 문서",
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading("1.4 표기 규칙", level=2)
dc.add_table(doc, ["표기", "의미"], [
    ["#nnn", "MAVLink 메시지 ID"],
    ["TBD", "본 초안에서 잠정 지정한 값으로, 배포 전 관계자 협의를 통해 최종 확정 필요"],
    ["float / int / uint", "IEEE-754 float32, 부호있는/없는 정수 (비트폭은 각 표에 명시)"],
    ["NaN", "Not-a-Number — 해당 필드가 무효(invalid)함을 의미 (PX4 관례)"],
], col_widths_cm=[3.5, 12.5])

# =====================================================================
doc.add_heading("2. 통신 프로토콜 개요", level=1)

doc.add_heading("2.1 MAVLink v2 프레임 구조", level=2)
dc.add_table(doc, ["필드", "크기", "설명"], [
    ["Magic (STX)", "1 byte", "0xFD (MAVLink v2 식별자)"],
    ["Len", "1 byte", "페이로드 길이"],
    ["Incompat/Compat Flags", "각 1 byte", "비호환/호환 플래그 (서명 사용 시 0x01)"],
    ["Seq", "1 byte", "패킷 순번 (0~255 롤오버, 유실 검출용)"],
    ["SysID / CompID", "각 1 byte", "송신 시스템/컴포넌트 ID"],
    ["Msg ID", "3 bytes", "메시지 ID (0~16777215, 커스텀 메시지는 통상 180 이상 대역 사용)"],
    ["Payload", "가변", "메시지별 정의된 필드"],
    ["Checksum", "2 bytes", "CRC-16/MCRF4XX"],
    ["Signature (옵션)", "13 bytes", "본 프로젝트에서는 폐쇄망 실험실 환경으로 미사용을 기본값으로 함"],
], col_widths_cm=[4.5, 2.5, 9.0])

doc.add_heading("2.2 시스템/컴포넌트 ID 할당", level=2)
dc.add_table(doc, ["대상", "System ID", "Component ID", "비고"], [
    ["FCC (PX4 OFP)", "1", "1 (MAV_COMP_ID_AUTOPILOT1)", "표준 PX4 기본값"],
    ["HILS 환경 시뮬레이터(ENV)", "1 (동일 System)", "26 (MAV_COMP_ID_ONBOARD_COMPUTER 대역, TBD)", "SITL 시뮬레이터 관례상 동일 System ID 공유"],
    ["시각화 시스템(VIS)", "245", "190 (MAV_COMP_ID_MISSIONPLANNER 대역, TBD)", "지상시스템(GCS) 대역 사용, 실제 값은 사내 MAVLink ID 대장에서 중복 확인 후 확정"],
], col_widths_cm=[5.0, 2.5, 5.5, 3.0])

doc.add_heading("2.3 논리 채널 요약", level=2)
dc.add_table(doc, ["채널", "전송 주체", "수신 주체", "핵심 메시지"], [
    ["A. 센서 주입", "ENV", "FCC", "HIL_SENSOR(#107), HIL_GPS(#113), HIL_TILT_STATE(TBD, 신규)"],
    ["B. 액추에이터/틸트 명령", "FCC", "ENV", "HIL_ACTUATOR_CONTROLS(#93), HIL_TILT_ACTUATOR_CONTROLS(TBD, 신규)"],
    ["C. 상태 텔레메트리", "FCC", "VIS", "GLOBAL_POSITION_INT(#33), LOCAL_POSITION_NED(#32), ATTITUDE(#30), VFR_HUD(#74), WIND_COV(#231), HIL_TILT_STATE(재송출)"],
    ["D. 조종입력 검증", "RC/GCS", "FCC", "RC_CHANNELS(#65), MANUAL_CONTROL(#69, 대안)"],
], col_widths_cm=[4.5, 2.5, 2.5, 6.5])

dc.add_master_parameter_list(doc, heading_text="2.4 ICD 파라미터 목록 (통합)", heading_level=2)

doc.add_paragraph(
    "위 통합 목록의 각 신호는 3장(표준 메시지) 및 4장(커스텀 메시지)의 프레임 구조 안에서 어느 필드에 "
    "해당하는지 상세 정의되어 있다. 실제 개발 시에는 본 목록을 기준으로 ENV/VIS/조종기 담당팀과 신호별 "
    "범위·비트폭·주기를 최종 확정한 뒤, 3~4장의 필드 표를 개정한다."
)

# =====================================================================
doc.add_heading("3. 표준 MAVLink 메시지 정의", level=1)

doc.add_heading("3.1 HIL_SENSOR (#107) — ENV → FCC", level=2)
doc.add_paragraph("가상 IMU/기압/지자기 센서 데이터를 FCC EKF2에 주입한다.")
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "센서 샘플 타임스탬프"],
    ["xacc, yacc, zacc", "float", "m/s²", "FRD 기체좌표계 가속도"],
    ["xgyro, ygyro, zgyro", "float", "rad/s", "FRD 기체좌표계 각속도"],
    ["xmag, ymag, zmag", "float", "gauss", "FRD 기체좌표계 지자기"],
    ["abs_pressure", "float", "hPa", "절대 기압"],
    ["diff_pressure", "float", "hPa", "차압(피토관, 대기속도 산출용)"],
    ["pressure_alt", "float", "m", "기압고도"],
    ["temperature", "float", "°C", "센서 온도"],
    ["fields_updated", "uint32", "bitmask", "이번 패킷에서 갱신된 필드 비트마스크"],
    ["id", "uint8", "-", "센서 인스턴스 ID (다중 IMU 지원 시)"],
], col_widths_cm=[3.5, 2.0, 2.0, 8.5])

doc.add_heading("3.2 HIL_GPS (#113) — ENV → FCC", level=2)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["fix_type", "uint8", "-", "0/1: no fix, 2: 2D, 3: 3D, 4~6: RTK 등 (SensorGps.FIX_TYPE_* 매핑)"],
    ["lat, lon", "int32", "degE7", "위도/경도 ×1e7"],
    ["alt", "int32", "mm", "AMSL 고도"],
    ["eph, epv", "uint16", "cm", "수평/수직 위치 정밀도"],
    ["vel", "uint16", "cm/s", "지상속도(스칼라)"],
    ["vn, ve, vd", "int16", "cm/s", "NED 속도 성분"],
    ["cog", "uint16", "cdeg", "이동방향(Course Over Ground) ×100"],
    ["satellites_visible", "uint8", "-", "가시 위성 수"],
    ["id", "uint8", "-", "GPS 인스턴스 ID"],
    ["yaw", "uint16", "cdeg", "듀얼안테나 헤딩(옵션), 0 = 미지원"],
], col_widths_cm=[3.5, 2.0, 2.0, 8.5])

doc.add_heading("3.3 HIL_STATE_QUATERNION (#115) — ENV → FCC (옵션/로깅용)", level=2)
doc.add_paragraph("제어 루프에는 필수가 아니며, 진리값(ground truth) 로깅·검증 목적으로만 사용을 권장한다.")
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["attitude_quaternion[4]", "float[4]", "-", "자세 쿼터니언 (w,x,y,z)"],
    ["rollspeed/pitchspeed/yawspeed", "float", "rad/s", "기체좌표계 각속도"],
    ["lat, lon, alt", "int32", "degE7 / mm", "위치(WGS84)"],
    ["vx, vy, vz", "int16", "cm/s", "NED 속도"],
    ["ind_airspeed, true_airspeed", "uint16", "cm/s", "지시/진대기속도"],
    ["xacc, yacc, zacc", "int16", "mG (1/1000 g)", "기체좌표계 가속도"],
], col_widths_cm=[4.5, 2.5, 3.0, 6.0])

doc.add_heading("3.4 HIL_ACTUATOR_CONTROLS (#93) — FCC → ENV", level=2)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["flags", "uint64", "bitmask", "제어모드 플래그(0: 기본)"],
    ["controls[16]", "float[16]", "정규화 [-1, 1] / NaN=disarm", "채널 0~3: 모터1~4 추력, 채널 4~7: 예비, 상세 채널맵은 4.3절 매핑표 참조"],
    ["mode", "uint8", "-", "MAV_MODE_FLAG 기반 armed 상태 등"],
], col_widths_cm=[3.5, 2.5, 4.0, 6.0])
dc.note(doc, "본 커스텀 기체는 모터 4식 + 틸트 서보 4식으로 구성되나, 틸트 명령은 정밀도 및 가독성을 위해 3.4절 채널을 재사용하지 않고 4.1절의 전용 커스텀 메시지(HIL_TILT_ACTUATOR_CONTROLS)로 분리 정의한다 (기획 확정 사항).")

doc.add_heading("3.5 텔레메트리 메시지 (FCC → VIS, EICD-02)", level=2)
dc.add_table(doc, ["메시지", "핵심 필드", "단위", "관련 uORB 소스"], [
    ["GLOBAL_POSITION_INT (#33)", "lat, lon, alt, relative_alt, vx, vy, vz, hdg", "degE7, mm, cm/s, cdeg", "VehicleGlobalPosition, VehicleLocalPosition"],
    ["LOCAL_POSITION_NED (#32)", "x, y, z, vx, vy, vz", "m, m/s", "VehicleLocalPosition"],
    ["ATTITUDE (#30)", "roll, pitch, yaw, rollspeed, pitchspeed, yawspeed", "rad, rad/s", "VehicleAttitude(쿼터니언→오일러 변환), VehicleAngularVelocity"],
    ["VFR_HUD (#74)", "airspeed, groundspeed, heading, throttle, alt, climb", "m/s, deg, %, m, m/s", "AirspeedValidated, VehicleGlobalPosition, VehicleLocalPosition"],
    ["WIND_COV (#231)", "wind_x, wind_y, wind_z, var_horiz, var_vert", "m/s", "Wind"],
], col_widths_cm=[3.5, 6.0, 2.5, 4.0])

# =====================================================================
doc.add_heading("4. 커스텀 메시지 정의 (신규)", level=1)
doc.add_paragraph(
    "MAVLink 표준셋에는 다중 틸트 서보의 각도/피드백을 표현하는 메시지가 없으므로, 본 프로젝트 전용 "
    "커스텀 메시지 2종을 신설한다. 메시지 ID는 사내 MAVLink dialect(XML) 확장 파일에 정의하며, 아래 "
    "ID는 초안 단계의 잠정값(TBD)이다. 배포 전 타 프로젝트/외부 GCS와의 ID 충돌 여부를 반드시 재확인한다."
)

doc.add_heading("4.1 HIL_TILT_ACTUATOR_CONTROLS (ID: 52000, TBD) — FCC → ENV", level=2)
doc.add_paragraph("FCC의 틸트 제어 모듈(tv_att_control)이 산출한 4개 서보 목표각 명령을 환경 시뮬레이터에 전달한다.")
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["tilt_fl", "float", "deg", "전좌(Front-Left) 서보 목표각 (0°=수직/멀티콥터, 90°=수평/고정익)"],
    ["tilt_fr", "float", "deg", "전우(Front-Right) 서보 목표각"],
    ["tilt_rl", "float", "deg", "후좌(Rear-Left) 서보 목표각"],
    ["tilt_rr", "float", "deg", "후우(Rear-Right) 서보 목표각"],
    ["collective_tilt_norm", "float", "정규화 [0,1]", "집단 틸트 정규값 (0=수직, 1=수평) — TiltrotorExtraControls 대응"],
    ["armed", "uint8", "bool", "서보 구동 활성화 여부"],
], col_widths_cm=[4.0, 2.5, 3.0, 6.5])
doc.add_paragraph("근거(uORB 원본 필드) — msg/TiltAngleSetpoint.msg, msg/TiltrotorExtraControls.msg:", style='Intense Quote')
dc.add_table(doc, ["uORB 메시지", "필드"], [
    ["TiltAngleSetpoint", "tilt_fl, tilt_fr, tilt_rl, tilt_rr (deg)"],
    ["TiltrotorExtraControls", "collective_tilt_normalized_setpoint [0,1], collective_thrust_normalized_setpoint [0,1]"],
], col_widths_cm=[5.0, 11.0])

doc.add_heading("4.2 HIL_TILT_STATE (ID: 52001, TBD) — ENV → FCC", level=2)
doc.add_paragraph(
    "환경 시뮬레이터가 서보 동특성 모델(응답지연, 속도제한, 전류/온도 모사)로 계산한 결과를 실제 "
    "Dynamixel 피드백과 동일한 형식으로 FCC에 되돌려준다. FCC 내부에서는 이 메시지를 수신하여 실기체와 "
    "동일하게 TiltStatus uORB 토픽으로 재발행하는 HIL 브리지 로직이 필요하다(§4.4 참조)."
)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["angle[4]", "float[4]", "deg", "서보 실제 각도 (순서: FL, FR, RL, RR)"],
    ["angular_velocity[4]", "float[4]", "deg/s", "서보 각속도"],
    ["current[4]", "int16[4]", "mA", "서보 구동 전류"],
    ["temperature[4]", "float[4]", "°C", "서보 온도"],
    ["voltage[4]", "float[4]", "V", "서보 입력 전압"],
    ["moving[4]", "uint8[4]", "bool", "서보 이동 중 여부"],
    ["comm_warn[4]", "uint8[4]", "bool", "통신 경고 플래그"],
    ["comm_critical[4]", "uint8[4]", "bool", "통신 심각 오류 플래그"],
    ["hardware_error_status[4]", "int8[4]", "bitmask", "하드웨어 오류 상태 (과열/과전류 등)"],
], col_widths_cm=[4.5, 2.5, 2.5, 6.5])
doc.add_paragraph("근거(uORB 원본) — msg/TiltStatus.msg 필드를 1:1로 매핑함 (필드명·배열순서 동일하게 유지하여 FCC 내부 재발행 로직을 단순화):", style='Intense Quote')
dc.add_table(doc, ["TiltStatus.msg 필드", "타입"], [
    ["goal_angle[4], angle[4], angular_velocity[4]", "float32[4]"],
    ["current[4]", "int16[4]"],
    ["voltage[4], temperature[4]", "float32[4]"],
    ["moving[4]", "uint8[4]"],
    ["packet_error_status[4], hardware_error_status[4]", "uint8[4] / int8[4]"],
    ["bus_watchdog[4]", "int8[4]"],
    ["comm_loss_count_warning[4], comm_loss_count_critical[4]", "uint8[4]"],
    ["comm_warn[4], comm_critical[4]", "bool[4]"],
], col_widths_cm=[10.0, 6.0])

doc.add_heading("4.3 HIL_ACTUATOR_CONTROLS 채널 매핑 (모터, 참고)", level=2)
dc.add_table(doc, ["controls[] 인덱스", "매핑 대상", "비고"], [
    ["0", "모터 1 (Front-Left)", "ActuatorMotors.control[0] 대응"],
    ["1", "모터 2 (Front-Right)", "ActuatorMotors.control[1] 대응"],
    ["2", "모터 3 (Rear-Left)", "ActuatorMotors.control[2] 대응"],
    ["3", "모터 4 (Rear-Right)", "ActuatorMotors.control[3] 대응"],
    ["4~15", "미사용(예비)", "향후 페이로드/추가 액추에이터 확장 여지"],
], col_widths_cm=[4.0, 6.0, 6.0])

doc.add_heading("4.4 FCC 내부 처리 요구사항 (구현 가이드)", level=2)
for t in [
    "HIL_TILT_ACTUATOR_CONTROLS 송신: tv_att_control이 발행하는 TiltAngleSetpoint/TiltrotorExtraControls uORB를 구독하여 HIL 모드일 때 동일 주기로 MAVLink 송신하는 브리지 로직을 simulator_mavlink 모듈에 추가.",
    "HIL_TILT_STATE 수신: SimulatorMavlink::handle_message_hil_tilt_state()(신규)를 추가하여 수신 필드를 TiltStatus uORB로 발행 — 기존 handle_message_hil_gps()와 동일한 패턴 적용.",
    "HIL 모드에서는 dynamixel_uart 실물 드라이버 퍼블리시를 비활성화하여 TiltStatus 토픽에 대한 발행자 중복(더블 퍼블리셔)을 방지할 것.",
]:
    doc.add_paragraph(t, style='List Bullet')

# =====================================================================
doc.add_heading("5. uORB ↔ MAVLink 필드 매핑 총괄표", level=1)

doc.add_heading("5.1 채널 A/B (HILS 폐루프)", level=2)
dc.add_table(doc, ["MAVLink 메시지", "방향", "대응 uORB 토픽"], [
    ["HIL_SENSOR", "ENV→FCC", "SensorCombined (accelerometer_m_s2, gyro_rad), sensor_baro, sensor_mag (내부 드라이버 등가 주입)"],
    ["HIL_GPS", "ENV→FCC", "SensorGps"],
    ["HIL_ACTUATOR_CONTROLS", "FCC→ENV", "ActuatorMotors.control[0..3]"],
    ["HIL_TILT_ACTUATOR_CONTROLS (신규)", "FCC→ENV", "TiltAngleSetpoint, TiltrotorExtraControls"],
    ["HIL_TILT_STATE (신규)", "ENV→FCC", "TiltStatus"],
], col_widths_cm=[5.5, 2.5, 8.0])

doc.add_heading("5.2 채널 C (시각화 텔레메트리)", level=2)
dc.add_table(doc, ["MAVLink 메시지", "대응 uORB 토픽", "핵심 필드"], [
    ["GLOBAL_POSITION_INT", "VehicleGlobalPosition", "lat, lon, alt"],
    ["LOCAL_POSITION_NED", "VehicleLocalPosition", "x, y, z, vx, vy, vz, heading"],
    ["ATTITUDE / ATTITUDE_QUATERNION", "VehicleAttitude", "q[4] (w,x,y,z)"],
    ["VFR_HUD", "AirspeedValidated, VehicleLocalPosition", "true_airspeed_m_s, heading, vz"],
    ["WIND_COV", "Wind", "windspeed_north, windspeed_east"],
    ["HIL_TILT_STATE (재송출)", "TiltStatus", "angle[4], angular_velocity[4]"],
    ["BATTERY_STATUS (옵션)", "BatteryStatus", "voltage_v, remaining"],
    ["DISTANCE_SENSOR (옵션)", "DistanceSensor", "current_distance"],
], col_widths_cm=[5.0, 5.5, 5.5])

# =====================================================================
doc.add_heading("6. 타이밍 및 성능 요구사항", level=1)
dc.add_table(doc, ["항목", "요구값", "비고"], [
    ["채널 A→FCC 처리 지연", "≤ 10 ms (수신~EKF2 반영)", "EKF2 지연보상 한계 내"],
    ["채널 B FCC→ENV 지연", "≤ 10 ms", "FDM 적분 스텝과 정합 필요"],
    ["채널 A/B 왕복 지연(RTT)", "≤ 20 ms", "EICD-01 3.1.4절과 동일"],
    ["채널 C 갱신 지연(FCC→VIS)", "≤ 50 ms", "시각적 체감 지연 최소화 목적"],
    ["지터(jitter) 허용치", "≤ 5 ms (표준편차 기준)", "고정 주기 발행기(hrt) 사용 권장"],
    ["패킷 유실 허용치", "채널 A/B: 0% (재전송/보간 불가 가정), 채널 C: ≤1%(뷰어 보간으로 흡수)", ""],
], col_widths_cm=[5.0, 6.0, 5.0])

# =====================================================================
doc.add_heading("7. 오류 처리", level=1)
dc.add_table(doc, ["오류 상황", "검출 방법", "처리 방안"], [
    ["HEARTBEAT 두절", "1.5초 이상 미수신", "5.3절(EICD) 페일세이프 절차 적용, 링크 상태를 UI에 경고 표시"],
    ["체크섬 오류", "MAVLink CRC 불일치", "해당 패킷 폐기, 시퀀스 갭 카운터 증가"],
    ["시퀀스 갭", "seq 필드 불연속", "유실률 통계에 반영, 임계치 초과 시 경고 로그"],
    ["HIL_TILT_STATE 필드 NaN", "필드값 검증", "직전 유효값 유지(hold-last) 또는 comm_critical 플래그 설정"],
    ["메시지 ID 충돌", "사내 dialect 등록 대장 대조", "배포 전 정적 검증(§4 TBD ID 확정 절차)으로 사전 예방"],
], col_widths_cm=[4.0, 5.5, 6.5])

# =====================================================================
doc.add_heading("8. 연동 시퀀스 (개요)", level=1)
for t in [
    "① [초기화] FCC 부팅 → SYS_HITL=1 파라미터 로드 → MAVLink 링크 오픈(채널 A/B), 별도 인스턴스로 채널 C 오픈",
    "② [연결수립] ENV/VIS ↔ FCC 간 HEARTBEAT 상호 교환, 버전/역량(capability) 확인",
    "③ [HIL 모드 진입] ENV가 HIL_SENSOR/HIL_GPS/HIL_TILT_STATE 주기 송신 시작 → FCC EKF2 정렬(수 초 내)",
    "④ [정상 운용] 채널 A/B 폐루프 지속(§6 타이밍 준수), 채널 C로 VIS에 상태 지속 송출",
    "⑤ [시나리오 종료] Disarm → SYS_HITL=0 복귀 또는 링크 종료(FIN 절차 없음, MAVLink 특성상 타임아웃 기반 종료)",
]:
    doc.add_paragraph(t, style='List Number')

# =====================================================================
doc.add_heading("9. 부록 A — 관련 uORB 메시지 원본 필드 발췌", level=1)
doc.add_paragraph("아래는 본 ICD의 근거가 된 저장소 내 msg/*.msg 정의에서 발췌한 핵심 필드이다(주석/일부 예약 필드 생략).")

doc.add_heading("A.1 VehicleGlobalPosition.msg", level=3)
dc.add_table(doc, ["필드", "타입", "단위"], [
    ["lat, lon", "float64", "degrees"],
    ["alt", "float32", "meters (AMSL)"],
    ["eph, epv", "float32", "meters"],
    ["terrain_alt", "float32", "meters (WGS84)"],
], col_widths_cm=[6.0, 4.0, 6.0])

doc.add_heading("A.2 VehicleLocalPosition.msg (발췌)", level=3)
dc.add_table(doc, ["필드", "타입", "단위"], [
    ["x, y, z", "float32", "meters (NED)"],
    ["vx, vy, vz", "float32", "m/s (NED)"],
    ["ax, ay, az", "float32", "m/s² (NED)"],
    ["heading", "float32", "rad (-PI..PI)"],
    ["dist_bottom", "float32", "meters"],
], col_widths_cm=[6.0, 4.0, 6.0])

doc.add_heading("A.3 VehicleAttitude.msg", level=3)
dc.add_table(doc, ["필드", "타입", "단위"], [
    ["q[4]", "float32[4]", "쿼터니언 (w,x,y,z), Hamilton, FRD→NED"],
], col_widths_cm=[6.0, 4.0, 6.0])

doc.add_heading("A.4 TiltStatus.msg (전체)", level=3)
dc.add_table(doc, ["필드", "타입", "단위/비고"], [
    ["packet_error_status[4]", "uint8[4]", "통신 패킷 오류 상태"],
    ["hardware_error_status[4]", "uint8[4]", "하드웨어 오류 상태"],
    ["bus_watchdog[4]", "int8[4]", "버스 워치독"],
    ["goal_angle[4]", "float32[4]", "목표각 (deg)"],
    ["realtime_tick[4]", "uint16[4]", "서보 내부 틱"],
    ["moving[4]", "uint8[4]", "이동 중 여부"],
    ["current[4]", "int16[4]", "전류 (mA)"],
    ["angular_velocity[4]", "float32[4]", "각속도 (deg/s)"],
    ["angle[4]", "float32[4]", "실제 각도 (deg)"],
    ["voltage[4]", "float32[4]", "전압 (V)"],
    ["temperature[4]", "float32[4]", "온도 (°C)"],
    ["comm_loss_count_warning[4] / _critical[4]", "uint8[4]", "통신 유실 카운터"],
    ["comm_warn[4] / comm_critical[4]", "bool[4]", "통신 상태 플래그"],
], col_widths_cm=[7.0, 3.5, 5.5])

doc.save("ICD-PXTR-HILS-001.docx")
print("ICD saved")
