# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import doc_common as dc

doc = Document()
dc.set_base_style(doc)
dc.add_page_number_footer(doc)
dc.enable_auto_update_fields(doc)

dc.add_title_page(
    doc,
    "PX4 기반 틸트로터 VTOL 비행제어시스템",
    "HILS/시각화 연동 인터페이스 통제 문서 통합본\n(EICD + ICD + 비행데이터 분석 파라미터 가이드)",
    doc_no="ICD-PXTR-HILS-000 (통합본)",
    rev="A (초안)",
    date_str="2026-08-31",
)

dc.add_revision_history(doc, [
    ["A", "2026-08-31", "최초 발행 — EICD-PXTR-HILS-001, ICD-PXTR-HILS-001 통합 및\n"
                          "Part III 비행데이터 분석 파라미터 가이드 신규 작성", "비행제어SW팀"],
])

doc.add_heading("통합본 안내", level=1)
doc.add_paragraph(
    "본 문서는 아래 3개 문서를 하나로 묶은 통합본이다. 개별 문서를 형상관리 단위로 별도 유지하는 경우 "
    "본 통합본은 해당 문서들의 스냅샷(참조용)으로 취급하며, 실제 개정은 각 개별 문서에서 수행한 뒤 "
    "본 통합본에 반영한다."
)
dc.add_table(doc, ["구성", "원본 문서", "핵심 질문에 대한 답"], [
    ["PART Ⅰ", "EICD-PXTR-HILS-001\n(외부 인터페이스 통제 문서)", "FCC · HILS 환경 시뮬레이터 · 시각화 시스템이\n무엇을 통해, 어떤 방향으로 연결되는가?"],
    ["PART Ⅱ", "ICD-PXTR-HILS-001\n(인터페이스 통제 문서)", "각 신호가 어떤 주기·범위·분해능으로,\n어떤 메시지 필드에 실려 오가는가?"],
    ["PART Ⅲ", "(신규) 비행데이터 분석 파라미터 가이드", "비행 후 로그(ULog)를 분석할 때\n어떤 파라미터를 봐야 하는가?"],
], col_widths_cm=[2.0, 6.0, 8.0])
doc.add_paragraph(
    "PART Ⅲ은 PART Ⅰ·Ⅱ에서 정의한 인터페이스 신호(A-xx/B-xx/C-xx)가 실제 비행 후에는 PX4 ULog의 "
    "uORB 토픽으로 고스란히 기록된다는 점에 착안하여, 두 문서를 '설계 시점의 인터페이스 정의'에서 "
    "'비행 후 로그 분석'까지 한 줄로 잇는 것을 목적으로 한다."
)

doc.add_heading("전체 목차 (Table of Contents)", level=1)
dc.add_toc(doc)
doc.add_page_break()


def part_divider(doc, roman, title_kr, subtitle):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PART {roman}")
    run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_kr)
    run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = dc.NAVY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    doc.add_page_break()


# =====================================================================
# PART I — EICD
# =====================================================================
part_divider(doc, "Ⅰ", "외부 인터페이스 통제 문서 (EICD)",
             "FCC ↔ HILS 환경 시뮬레이터 ↔ 시각화 시스템 간 물리/논리 인터페이스 정의")

doc.add_heading("1. 개요", level=1)

doc.add_heading("1.1 목적", level=2)
doc.add_paragraph(
    "본 Part는 Airbility 커스텀 PX4 비행제어소프트웨어(OFP: Operational Flight Program)가 "
    "HILS(Hardware-In-the-Loop Simulation) 환경 시뮬레이터 및 시각화(Visualization) 시스템과 "
    "연동하기 위해 필요한 외부 인터페이스(물리적/논리적 연결, 프로토콜, 방향, 주기)를 정의하는 것을 "
    "목적으로 한다. 각 인터페이스에서 교환되는 메시지/필드 단위의 상세 정의는 PART Ⅱ(ICD)를 따른다."
)

doc.add_heading("1.2 적용범위", level=2)
doc.add_paragraph("본 Part는 아래 3개 서브시스템 간 외부 인터페이스에 적용된다.")
for t in [
    "① 비행제어시스템(FCC, Flight Control Computer) — PX4 커스텀 펌웨어 (airbility-dev/px4-custom-firmware, develop/main-v1.15.4 기준)",
    "② HILS 환경 시뮬레이터(ENV) — 비행 동역학(FDM) 및 센서모델을 계산하는 외부 소프트웨어/장비",
    "③ 시각화 시스템(VIS) — 조종사/관제사에게 3D 외부 시계(Out-The-Window) 또는 지도/HUD 화면을 제공하는 소프트웨어",
]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_paragraph(
    "실비행(운용) 모드에서 FCC와 실제 센서/액추에이터(IMU, GPS, 틸트 서보 등) 간 인터페이스는 "
    "본 Part 3.3절(참고용)로 다루되, 상세 하드웨어 ICD는 별도 문서로 관리한다."
)

doc.add_heading("1.3 용어 및 약어 정의", level=2)
dc.add_table(doc, ["약어", "정의"], [
    ["FCC", "Flight Control Computer — 비행제어컴퓨터 (본 프로젝트에서는 PX4 OFP가 탑재되는 보드)"],
    ["OFP", "Operational Flight Program — 비행제어소프트웨어(펌웨어)"],
    ["HILS / HITL", "Hardware-In-the-Loop Simulation — 실제 FCC 하드웨어에 가상 센서 입력을 주입하여 폐루프로 검증하는 시뮬레이션 기법"],
    ["SITL", "Software-In-the-Loop — FCC 소프트웨어를 PC 상에서 실행하는 순수 소프트웨어 시뮬레이션"],
    ["ENV", "HILS 환경 시뮬레이터 (비행 동역학 모델 + 센서 모델 + 지형/기상 모델)"],
    ["VIS", "시각화 시스템 (3D 뷰, 지도, HUD 등)"],
    ["EICD", "External Interface Control Document — 외부 인터페이스 통제 문서"],
    ["ICD", "Interface Control Document — 메시지/필드 단위 상세를 정의하는 인터페이스 통제 문서"],
    ["MAVLink", "PX4/ArduPilot 계열에서 표준으로 사용하는 경량 UAV 통신 프로토콜"],
    ["uORB", "PX4 내부 프로세스 간 발행-구독(pub-sub) 미들웨어 및 그 메시지(토픽)"],
    ["ULog", "PX4 비행 중 uORB 토픽을 그대로 기록하는 바이너리 로그 포맷(.ulg) — PART Ⅲ의 분석 대상"],
    ["EKF2", "PX4의 확장 칼만필터 기반 상태추정 모듈 (위치/속도/자세 추정)"],
    ["NED", "North-East-Down — 로컬 지평좌표계"],
    ["FRD", "Forward-Right-Down — 기체(Body) 좌표계"],
    ["WGS84", "World Geodetic System 1984 — GPS 위경도 기준 좌표계"],
    ["HIL_*", "MAVLink 표준 HIL(Hardware-In-the-Loop) 메시지 군 (HIL_SENSOR, HIL_GPS 등)"],
], col_widths_cm=[3.0, 13.0])

doc.add_heading("1.4 참고 문서", level=2)
for t in [
    "PX4 사용자/개발자 가이드 (https://docs.px4.io/main/en/)",
    "MAVLink Common/Development Message Set (https://mavlink.io/en/messages/)",
    "airbility-dev/px4-custom-firmware 소스코드 (msg/, src/modules/simulation/, src/modules/mavlink/)",
    "사내 ULog 분석 도구 ulog2csv.py / flight_plotter.py (PART Ⅲ에서 상세 연계)",
]:
    doc.add_paragraph(t, style='List Bullet')

# --- 2. 시스템 개요 ---
doc.add_heading("2. 시스템 개요", level=1)

doc.add_heading("2.1 시스템 구성", level=2)
doc.add_paragraph(
    "본 시스템은 실물 FCC 보드(PX4 커스텀 OFP 탑재)를 중심으로, HILS 환경 시뮬레이터가 비행 동역학과 "
    "센서 신호를 계산하여 FCC에 주입하고, FCC는 이를 실제 비행 상황과 동일하게 인식하여 제어연산을 수행한 뒤 "
    "액추에이터(모터/틸트 서보) 명령을 다시 환경 시뮬레이터로 되돌려 폐루프를 형성한다. 동시에 FCC가 "
    "추정한 상태값(위치/속도/자세/틸트각 등)은 시각화 시스템으로 전달되어 외부 환경(시점, 지형, HUD 등)을 "
    "실시간으로 갱신한다."
)
p = doc.add_paragraph()
run = p.add_run(
    "[FCC: PX4 OFP]\n"
    "   ↑ᴬ HIL 센서/GPS/틸트 피드백 주입 (ENV→FCC)\n"
    "   ↓ᴮ 액추에이터/틸트 명령 출력 (FCC→ENV)\n"
    "[ENV: HILS 환경 시뮬레이터 — FDM/센서모델/지형·기상모델]\n"
    "\n"
    "[FCC: PX4 OFP]\n"
    "   ↓ᶜ 상태 텔레메트리 (FCC→VIS)\n"
    "[VIS: 시각화 시스템 — 3D 뷰/지도/HUD]\n"
    "\n"
    "[FCC: PX4 OFP] --- 비행 중 전체 uORB 토픽 기록 ---> [ULog 파일(.ulg)] --- PART Ⅲ ---> [분석자]"
)
run.font.name = "Consolas"; run.font.size = Pt(9.5)

doc.add_heading("2.2 인터페이스 목록", level=2)
dc.add_table(doc, ["ID", "명칭", "연결 시스템", "매체", "방향", "용도"], [
    ["EICD-01", "FCC ↔ HILS 환경 시뮬레이터\n(센서 주입 / 액추에이터 피드백)", "FCC ↔ ENV", "Ethernet(UDP) 또는 USB 시리얼", "양방향", "HILS 폐루프 핵심 인터페이스"],
    ["EICD-02", "FCC → 시각화 시스템\n(상태 텔레메트리)", "FCC → VIS", "Ethernet(UDP/TCP) 또는 시리얼", "단방향\n(FCC→VIS)", "외부 환경/HUD 갱신용 상태 송출"],
    ["EICD-03", "FCC ↔ 틸트 액추에이터 하드웨어\n(Dynamixel, UART)", "FCC ↔ 틸트 서보 4식", "UART (반이중, RS-485/TTL)", "양방향", "실비행 시 실제 틸트 구동/피드백\n(HILS 모드에서는 EICD-01의\nHIL_TILT_* 메시지로 대체됨)"],
    ["EICD-04\n(참고)", "ENV ↔ VIS\n(환경/지형 상태 공유)", "ENV ↔ VIS", "구현 방식에 따름", "양방향", "본 EICD 범위 밖(ENV·VIS 내부 구현),\n필요 시 별도 ICD로 관리"],
], col_widths_cm=[1.8, 4.2, 3.0, 3.5, 2.0, 3.5])

# --- 3. 인터페이스별 상세 정의 ---
doc.add_heading("3. 인터페이스별 상세 정의", level=1)

doc.add_heading("3.1 EICD-01: FCC ↔ HILS 환경 시뮬레이터", level=2)

doc.add_heading("3.1.1 물리 인터페이스", level=3)
dc.add_table(doc, ["항목", "내용"], [
    ["연결 매체", "1순위: Ethernet (UDP/IP) — 개발/실험실 환경 권장\n2순위: USB-시리얼 (FCC의 TELEM 포트) — 실기체 텔레메트리 포트를 그대로 사용할 경우"],
    ["네트워크 포트(UDP 예시)", "FCC 수신(센서 주입): UDP 4560\nFCC 송신(액추에이터 피드백): UDP 14580\n(PX4 SITL/HITL 기본 포트 관례를 계승, 실제 값은 사내 네트워크 설계서에서 최종 확정)"],
    ["시리얼 파라미터(대안)", "Baudrate 921600 bps, 8N1, HW 흐름제어 없음"],
    ["케이블/커넥터", "Ethernet: RJ45 / USB-시리얼: Micro-USB 또는 JST-GH (보드 사양에 따름)"],
], col_widths_cm=[4.0, 12.0])

doc.add_heading("3.1.2 논리 인터페이스 (프로토콜)", level=3)
doc.add_paragraph(
    "본 인터페이스는 PX4가 표준으로 지원하는 MAVLink v2 HIL 메시지 세트를 기반으로 하며, "
    "PX4 소스 내 simulator_mavlink 모듈(src/modules/simulation/simulator_mavlink/SimulatorMavlink.cpp)에서 "
    "그대로 처리 가능한 방식을 채택한다."
)
dc.add_table(doc, ["방향", "논리 채널", "대표 메시지", "설명"], [
    ["ENV → FCC", "센서 주입", "HIL_SENSOR, HIL_GPS,\n(옵션) HIL_OPTICAL_FLOW,\nDISTANCE_SENSOR", "IMU/GPS/기압/지자기 등 가상 센서값을 FCC의 EKF2에 주입"],
    ["ENV → FCC", "틸트 피드백 (커스텀)", "HIL_TILT_STATE (신규)", "4개 틸트 서보의 실제 각도/각속도/전류/온도 피드백을 시뮬레이션 값으로 대체 주입"],
    ["FCC → ENV", "액추에이터 출력", "HIL_ACTUATOR_CONTROLS", "모터 추력 명령(4식) — 환경 시뮬레이터의 FDM 입력으로 사용"],
    ["FCC → ENV", "틸트 명령 (커스텀)", "HIL_TILT_ACTUATOR_CONTROLS (신규)", "4개 틸트 서보 목표각(setpoint) 명령"],
    ["FCC → ENV", "상태 비교/로깅(옵션)", "HIL_STATE_QUATERNION", "FCC 자체 추정치와 진리값(ground truth) 비교용"],
], col_widths_cm=[2.2, 2.8, 4.0, 7.0])
dc.note(doc, "HIL_TILT_STATE, HIL_TILT_ACTUATOR_CONTROLS는 MAVLink 표준셋에 없는 신규 메시지로, 상세 정의 및 임시 메시지 ID는 PART Ⅱ 4장을 따른다.")

doc.add_heading("3.1.3 데이터 갱신 주기 요구사항", level=3)
dc.add_table(doc, ["데이터", "최소 주기", "권장 주기", "비고"], [
    ["IMU (가속도/각속도)", "150 Hz", "250 Hz", "EKF2 예측 스텝과 정합되어야 함"],
    ["GPS", "5 Hz", "10 Hz", "실제 GPS 모듈 갱신율과 동일하게 설정"],
    ["기압/자기장", "10 Hz", "50 Hz", ""],
    ["틸트 피드백 (HIL_TILT_STATE)", "50 Hz", "100 Hz", "Dynamixel 실제 피드백 주기(TiltStatus)와 정합"],
    ["액추에이터 출력 (FCC→ENV)", "FCC 제어루프 주기와 동일", "250~400 Hz", "mc/tv 제어 루프 주기에 종속"],
], col_widths_cm=[5.0, 3.0, 3.0, 5.0])

doc.add_heading("3.1.4 동기화 및 시간 기준", level=3)
doc.add_paragraph(
    "모든 메시지는 마이크로초(usec) 단위의 단조증가 타임스탬프를 포함한다. FCC와 ENV는 세션 시작 시 "
    "하트비트(HEARTBEAT) 교환을 통해 시각 기준(epoch)을 정렬하며, 왕복 지연(RTT)은 20 ms 이내를 목표로 한다."
)

doc.add_heading("3.2 EICD-02: FCC → 시각화 시스템", level=2)
dc.add_table(doc, ["항목", "내용"], [
    ["물리 인터페이스", "Ethernet(UDP/TCP), 필요 시 시리얼 텔레메트리 포트 분기"],
    ["논리 프로토콜", "MAVLink v2 텔레메트리 스트림 (GLOBAL_POSITION_INT, LOCAL_POSITION_NED, ATTITUDE, VFR_HUD, WIND_COV) + 커스텀 HIL_TILT_STATE 재송출"],
    ["방향", "단방향 (FCC → VIS)"],
    ["권장 갱신 주기", "위치/속도: 20~50 Hz, 자세: 50~100 Hz, 틸트 상태: 20 Hz"],
    ["연결 방식", "FCC의 두 번째 MAVLink 인스턴스(-i 1) 또는 UDP 브로드캐스트로 EICD-01과 물리적으로 분리 권장"],
], col_widths_cm=[4.0, 12.0])

doc.add_heading("3.3 EICD-03: FCC ↔ 틸트 액추에이터 하드웨어 (참고, 실비행 기준)", level=2)
doc.add_paragraph(
    "실비행(실기체) 구성에서의 물리 인터페이스이며, HILS 모드에서는 본 채널이 비활성화되고 동일한 역할을 "
    "EICD-01의 HIL_TILT_ACTUATOR_CONTROLS / HIL_TILT_STATE 메시지가 대신한다."
)
dc.add_table(doc, ["항목", "내용"], [
    ["매체", "UART (dynamixel_uart, uart_reader 모듈에서 처리)"],
    ["대상 장치", "틸트 서보 4식 (FL/FR/RL/RR, Dynamixel 프로토콜)"],
    ["FCC → 서보", "목표각 명령 (TiltAngleSetpoint uORB → Dynamixel Goal Position)"],
    ["서보 → FCC", "각도/각속도/전류/온도/통신상태 피드백 (TiltStatus uORB로 발행)"],
], col_widths_cm=[4.0, 12.0])

# --- 4. 좌표계 ---
doc.add_heading("4. 좌표계, 단위 및 시간 기준", level=1)
dc.add_table(doc, ["항목", "정의"], [
    ["위치(전역)", "WGS84 위도/경도(도, deg), 고도는 AMSL 기준(m)"],
    ["위치(로컬)", "NED(North-East-Down) 좌표계, FCC 부팅/EKF2 초기화 시점을 원점으로 함 (단위 m)"],
    ["자세", "쿼터니언 q=(w,x,y,z), Hamilton 컨벤션, FRD(기체) → NED(지면) 회전으로 정의"],
    ["각속도", "FRD 기체좌표계 기준 rad/s"],
    ["속도", "NED 좌표계 기준 m/s"],
    ["틸트각", "서보 축 기준 deg (0°: 수직/멀티콥터 모드, 90°: 수평/고정익 모드)"],
    ["시간", "메시지 내 time_usec: 시스템 부팅 후 경과 마이크로초(단조증가)"],
], col_widths_cm=[3.5, 12.5])

# --- 5. 운용 모드 ---
doc.add_heading("5. 운용 모드 및 천이", level=1)
dc.add_table(doc, ["모드", "설명", "센서/액추에이터 경로"], [
    ["실비행 (FLIGHT)", "실제 기체에서 운용", "실 IMU/GPS/틸트 서보 (EICD-03 경로)"],
    ["SITL", "FCC 소프트웨어를 PC에서 실행", "PX4 내장 시뮬레이터(jMAVSim/Gazebo 등)"],
    ["HITL/HILS (본 문서 대상)", "실물 FCC 보드 + 외부 환경 시뮬레이터", "EICD-01 (HIL_* 메시지 경로)"],
], col_widths_cm=[3.5, 6.5, 6.0])

doc.add_heading("5.1 모드 진입 절차 (개요)", level=2)
for t in [
    "① FCC 파라미터 SYS_HITL = 1 설정",
    "② FCC ↔ ENV 간 MAVLink 링크 연결 및 HEARTBEAT 교환 확인",
    "③ ENV로부터 HIL_SENSOR/HIL_GPS 수신 시작 확인 (EKF2 정렬 여부)",
    "④ 틸트 인터페이스 정상 여부 확인 (HIL_TILT_STATE → TiltStatus uORB 확인)",
    "⑤ 안전 절차에 따라 Arming 후 HILS 시나리오 시작",
]:
    doc.add_paragraph(t, style='List Number')

doc.add_heading("5.2 페일세이프", level=2)
dc.add_table(doc, ["이상 상황", "판단 기준", "FCC 동작"], [
    ["ENV 링크 두절", "HEARTBEAT 미수신 1.5초 이상", "HITL 데이터 소스 무효 처리, 경고, (설정 시) 강제 Disarm"],
    ["센서 값 미갱신", "HIL_SENSOR 미수신 500ms 이상", "EKF2 타임아웃 처리, 안전상태 전이"],
    ["틸트 피드백 두절", "HIL_TILT_STATE 미수신 300ms 이상", "comm_loss 카운터 증가 → 안전모드 전이 검토"],
], col_widths_cm=[4.0, 5.5, 6.5])

# --- 6. 검증 ---
doc.add_heading("6. 인터페이스 검증 및 형상관리", level=1)
for t in [
    "루프백 테스트: 더미 송신기로 HIL_SENSOR/HIL_GPS를 주입하여 FCC 로그(EKF2 status)로 정상 소비 여부 확인",
    "필드 단위 검증: uORB 리스너(listener 커맨드)로 스케일/단위/부호 일치 확인",
    "폐루프 검증: 계단입력(step input)에 대한 상호 발산 여부 확인",
    "틸트 인터페이스 검증: HIL_TILT_STATE에 알려진 각도를 주입하고 TiltStatus 및 tv_att_control 반응 확인",
]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_paragraph(
    "본 문서는 비행제어SW팀이 형상관리 주체이며, 인터페이스 변경 시 PART Ⅱ와 동시 개정하고 개정이력에 반영한다."
)

# --- 7. 부록 ---
doc.add_heading("7. 부록: 관련 소스 모듈 매핑", level=1)
dc.add_table(doc, ["인터페이스", "관련 PX4 소스 모듈/경로"], [
    ["EICD-01 (HIL 센서/액추에이터)", "src/modules/simulation/simulator_mavlink/SimulatorMavlink.cpp"],
    ["EICD-01 (개별 센서 시뮬레이션 참고)", "src/modules/simulation/sensor_gps_sim, sensor_baro_sim, sensor_mag_sim, sensor_airspeed_sim, battery_simulator"],
    ["EICD-02 (텔레메트리)", "src/modules/mavlink/streams/*.hpp"],
    ["EICD-03 (실 틸트 서보)", "src/modules/dynamixel_uart/, src/modules/uart_reader/, msg/TiltStatus.msg, msg/TiltAngleSetpoint.msg"],
    ["상위 제어(틸트로터)", "src/modules/tv_att_control/, src/modules/tv_pos_control/, src/modules/tv_control_allocator/"],
    ["Gazebo 연동 참고", "src/modules/simulation/gz_bridge"],
], col_widths_cm=[4.5, 11.5])


# =====================================================================
# PART II — ICD
# =====================================================================
doc.add_page_break()
part_divider(doc, "Ⅱ", "인터페이스 통제 문서 (ICD)",
             "메시지/필드 단위 상세 — 신호별 전송경로·주기·범위·분해능")

doc.add_heading("8. ICD 개요", level=1)
doc.add_heading("8.1 목적", level=2)
doc.add_paragraph(
    "본 Part는 PART Ⅰ에서 정의한 외부 인터페이스(EICD-01, EICD-02)에 대해, 실제로 송수신되는 메시지의 "
    "프레임 구조, 필드, 자료형, 단위, 유효범위를 소프트웨어 구현 수준에서 정의한다."
)
doc.add_heading("8.2 적용범위", level=2)
dc.add_table(doc, ["채널", "방향", "목적", "관련 EICD"], [
    ["A. 센서/피드백 주입", "ENV → FCC", "가상 IMU/GPS/기압/자기장 및 틸트 서보 피드백을 FCC에 주입", "EICD-01"],
    ["B. 액추에이터/틸트 명령", "FCC → ENV", "모터 추력 및 틸트 목표각 명령을 ENV의 동역학 모델에 전달", "EICD-01"],
    ["C. 상태 텔레메트리", "FCC → VIS", "위치/속도/자세/풍향/틸트상태 등 시각화용 상태값 송출", "EICD-02"],
    ["D. 조종입력 검증", "RC/GCS → FCC", "조종기 입력을 FCC에 전달하여 조종 반응 검증", "EICD-05"],
], col_widths_cm=[4.0, 2.5, 6.5, 3.0])

doc.add_heading("8.3 참고 문서", level=2)
for t in [
    "MAVLink Common Message Set — https://mavlink.io/en/messages/common.html",
    "MAVLink Development Message Set — https://mavlink.io/en/messages/development.html",
    "PX4 uORB 메시지 정의 (airbility-dev/px4-custom-firmware, msg/*.msg)",
    "src/modules/simulation/simulator_mavlink/SimulatorMavlink.cpp",
    "PART Ⅰ 외부 인터페이스 통제 문서",
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading("8.4 표기 규칙", level=2)
dc.add_table(doc, ["표기", "의미"], [
    ["#nnn", "MAVLink 메시지 ID"],
    ["TBD", "본 초안에서 잠정 지정한 값으로, 배포 전 관계자 협의를 통해 최종 확정 필요"],
    ["float / int / uint", "IEEE-754 float32, 부호있는/없는 정수 (비트폭은 각 표에 명시)"],
    ["NaN", "Not-a-Number — 해당 필드가 무효(invalid)함을 의미 (PX4 관례)"],
], col_widths_cm=[3.5, 12.5])

doc.add_heading("9. 통신 프로토콜 개요", level=1)
doc.add_heading("9.1 MAVLink v2 프레임 구조", level=2)
dc.add_table(doc, ["필드", "크기", "설명"], [
    ["Magic (STX)", "1 byte", "0xFD (MAVLink v2 식별자)"],
    ["Len", "1 byte", "페이로드 길이"],
    ["Incompat/Compat Flags", "각 1 byte", "비호환/호환 플래그 (서명 사용 시 0x01)"],
    ["Seq", "1 byte", "패킷 순번 (0~255 롤오버, 유실 검출용)"],
    ["SysID / CompID", "각 1 byte", "송신 시스템/컴포넌트 ID"],
    ["Msg ID", "3 bytes", "메시지 ID (0~16777215, 커스텀 메시지는 통상 180 이상 대역 사용)"],
    ["Payload", "가변", "메시지별 정의된 필드"],
    ["Checksum", "2 bytes", "CRC-16/MCRF4XX"],
    ["Signature (옵션)", "13 bytes", "본 프로젝트에서는 폐쇄망 실험실 환경으로 미사용을 기본값으로 함"],
], col_widths_cm=[4.5, 2.5, 9.0])

doc.add_heading("9.2 시스템/컴포넌트 ID 할당", level=2)
dc.add_table(doc, ["대상", "System ID", "Component ID", "비고"], [
    ["FCC (PX4 OFP)", "1", "1 (MAV_COMP_ID_AUTOPILOT1)", "표준 PX4 기본값"],
    ["HILS 환경 시뮬레이터(ENV)", "1 (동일 System)", "26 (MAV_COMP_ID_ONBOARD_COMPUTER 대역, TBD)", "SITL 시뮬레이터 관례상 동일 System ID 공유"],
    ["시각화 시스템(VIS)", "245", "190 (MAV_COMP_ID_MISSIONPLANNER 대역, TBD)", "지상시스템(GCS) 대역 사용, 중복 확인 후 확정"],
], col_widths_cm=[5.0, 2.5, 5.5, 3.0])

doc.add_heading("9.3 논리 채널 요약", level=2)
dc.add_table(doc, ["채널", "전송 주체", "수신 주체", "핵심 메시지"], [
    ["A. 센서 주입", "ENV", "FCC", "HIL_SENSOR(#107), HIL_GPS(#113), HIL_TILT_STATE(TBD, 신규)"],
    ["B. 액추에이터/틸트 명령", "FCC", "ENV", "HIL_ACTUATOR_CONTROLS(#93), HIL_TILT_ACTUATOR_CONTROLS(TBD, 신규)"],
    ["C. 상태 텔레메트리", "FCC", "VIS", "GLOBAL_POSITION_INT(#33), LOCAL_POSITION_NED(#32), ATTITUDE(#30), VFR_HUD(#74), WIND_COV(#231), HIL_TILT_STATE(재송출)"],
    ["D. 조종입력 검증", "RC/GCS", "FCC", "RC_CHANNELS(#65), MANUAL_CONTROL(#69, 대안)"],
], col_widths_cm=[4.5, 2.5, 2.5, 6.5])

dc.add_master_parameter_list(doc, heading_text="9.4 ICD 파라미터 목록 (통합)", heading_level=2)

doc.add_paragraph(
    "위 통합 목록의 각 신호는 10장(표준 메시지) 및 11장(커스텀 메시지)의 프레임 구조 안에서 어느 필드에 "
    "해당하는지 상세 정의되어 있으며, PART Ⅲ에서는 이 신호들이 비행 후 어떤 uORB 로그 파라미터로 남는지 "
    "연결한다."
)

doc.add_heading("10. 표준 MAVLink 메시지 정의", level=1)

doc.add_heading("10.1 HIL_SENSOR (#107) — ENV → FCC", level=2)
doc.add_paragraph("가상 IMU/기압/지자기 센서 데이터를 FCC EKF2에 주입한다.")
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "센서 샘플 타임스탬프"],
    ["xacc, yacc, zacc", "float", "m/s²", "FRD 기체좌표계 가속도"],
    ["xgyro, ygyro, zgyro", "float", "rad/s", "FRD 기체좌표계 각속도"],
    ["xmag, ymag, zmag", "float", "gauss", "FRD 기체좌표계 지자기"],
    ["abs_pressure", "float", "hPa", "절대 기압"],
    ["diff_pressure", "float", "hPa", "차압(피토관, 대기속도 산출용)"],
    ["pressure_alt", "float", "m", "기압고도"],
    ["temperature", "float", "°C", "센서 온도"],
    ["fields_updated", "uint32", "bitmask", "이번 패킷에서 갱신된 필드 비트마스크"],
    ["id", "uint8", "-", "센서 인스턴스 ID (다중 IMU 지원 시)"],
], col_widths_cm=[3.5, 2.0, 2.0, 8.5])

doc.add_heading("10.2 HIL_GPS (#113) — ENV → FCC", level=2)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["fix_type", "uint8", "-", "0/1: no fix, 2: 2D, 3: 3D, 4~6: RTK 등"],
    ["lat, lon", "int32", "degE7", "위도/경도 ×1e7"],
    ["alt", "int32", "mm", "AMSL 고도"],
    ["eph, epv", "uint16", "cm", "수평/수직 위치 정밀도"],
    ["vel", "uint16", "cm/s", "지상속도(스칼라)"],
    ["vn, ve, vd", "int16", "cm/s", "NED 속도 성분"],
    ["cog", "uint16", "cdeg", "이동방향(Course Over Ground) ×100"],
    ["satellites_visible", "uint8", "-", "가시 위성 수"],
    ["id", "uint8", "-", "GPS 인스턴스 ID"],
    ["yaw", "uint16", "cdeg", "듀얼안테나 헤딩(옵션), 0 = 미지원"],
], col_widths_cm=[3.5, 2.0, 2.0, 8.5])

doc.add_heading("10.3 HIL_STATE_QUATERNION (#115) — ENV → FCC (옵션/로깅용)", level=2)
doc.add_paragraph("제어 루프에는 필수가 아니며, 진리값(ground truth) 로깅·검증 목적으로만 사용을 권장한다.")
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["attitude_quaternion[4]", "float[4]", "-", "자세 쿼터니언 (w,x,y,z)"],
    ["rollspeed/pitchspeed/yawspeed", "float", "rad/s", "기체좌표계 각속도"],
    ["lat, lon, alt", "int32", "degE7 / mm", "위치(WGS84)"],
    ["vx, vy, vz", "int16", "cm/s", "NED 속도"],
    ["ind_airspeed, true_airspeed", "uint16", "cm/s", "지시/진대기속도"],
    ["xacc, yacc, zacc", "int16", "mG (1/1000 g)", "기체좌표계 가속도"],
], col_widths_cm=[4.5, 2.5, 3.0, 6.0])

doc.add_heading("10.4 HIL_ACTUATOR_CONTROLS (#93) — FCC → ENV", level=2)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["flags", "uint64", "bitmask", "제어모드 플래그(0: 기본)"],
    ["controls[16]", "float[16]", "정규화 [-1, 1] / NaN=disarm", "채널 0~3: 모터1~4 추력, 채널 4~15: 예비"],
    ["mode", "uint8", "-", "MAV_MODE_FLAG 기반 armed 상태 등"],
], col_widths_cm=[3.5, 2.5, 4.0, 6.0])

doc.add_heading("10.5 텔레메트리 메시지 (FCC → VIS, EICD-02)", level=2)
dc.add_table(doc, ["메시지", "핵심 필드", "단위", "관련 uORB 소스"], [
    ["GLOBAL_POSITION_INT (#33)", "lat, lon, alt, relative_alt, vx, vy, vz, hdg", "degE7, mm, cm/s, cdeg", "VehicleGlobalPosition, VehicleLocalPosition"],
    ["LOCAL_POSITION_NED (#32)", "x, y, z, vx, vy, vz", "m, m/s", "VehicleLocalPosition"],
    ["ATTITUDE (#30)", "roll, pitch, yaw, rollspeed, pitchspeed, yawspeed", "rad, rad/s", "VehicleAttitude, VehicleAngularVelocity"],
    ["VFR_HUD (#74)", "airspeed, groundspeed, heading, throttle, alt, climb", "m/s, deg, %, m, m/s", "AirspeedValidated, VehicleGlobalPosition, VehicleLocalPosition"],
    ["WIND_COV (#231)", "wind_x, wind_y, wind_z, var_horiz, var_vert", "m/s", "Wind"],
], col_widths_cm=[3.5, 6.0, 2.5, 4.0])

doc.add_heading("11. 커스텀 메시지 정의 (신규)", level=1)
doc.add_paragraph(
    "MAVLink 표준셋에는 다중 틸트 서보의 각도/피드백을 표현하는 메시지가 없으므로, 본 프로젝트 전용 "
    "커스텀 메시지 2종을 신설한다. 아래 ID는 초안 단계의 잠정값(TBD)이다."
)

doc.add_heading("11.1 HIL_TILT_ACTUATOR_CONTROLS (ID: 52000, TBD) — FCC → ENV", level=2)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["tilt_fl", "float", "deg", "전좌(Front-Left) 서보 목표각"],
    ["tilt_fr", "float", "deg", "전우(Front-Right) 서보 목표각"],
    ["tilt_rl", "float", "deg", "후좌(Rear-Left) 서보 목표각"],
    ["tilt_rr", "float", "deg", "후우(Rear-Right) 서보 목표각"],
    ["collective_tilt_norm", "float", "정규화 [0,1]", "집단 틸트 정규값 (0=수직, 1=수평)"],
    ["armed", "uint8", "bool", "서보 구동 활성화 여부"],
], col_widths_cm=[4.0, 2.5, 3.0, 6.5])
doc.add_paragraph("근거(uORB 원본 필드) — msg/TiltAngleSetpoint.msg, msg/TiltrotorExtraControls.msg", style='Intense Quote')

doc.add_heading("11.2 HIL_TILT_STATE (ID: 52001, TBD) — ENV → FCC", level=2)
dc.add_table(doc, ["필드", "타입", "단위", "설명"], [
    ["time_usec", "uint64", "µs", "타임스탬프"],
    ["angle[4]", "float[4]", "deg", "서보 실제 각도 (순서: FL, FR, RL, RR)"],
    ["angular_velocity[4]", "float[4]", "deg/s", "서보 각속도"],
    ["current[4]", "int16[4]", "mA", "서보 구동 전류"],
    ["temperature[4]", "float[4]", "°C", "서보 온도"],
    ["voltage[4]", "float[4]", "V", "서보 입력 전압"],
    ["moving[4]", "uint8[4]", "bool", "서보 이동 중 여부"],
    ["comm_warn[4] / comm_critical[4]", "uint8[4]", "bool", "통신 경고/심각 오류 플래그"],
    ["hardware_error_status[4]", "int8[4]", "bitmask", "하드웨어 오류 상태"],
], col_widths_cm=[4.5, 2.5, 2.5, 6.5])
doc.add_paragraph("근거(uORB 원본) — msg/TiltStatus.msg 필드를 1:1로 매핑함", style='Intense Quote')

doc.add_heading("11.3 FCC 내부 처리 요구사항 (구현 가이드)", level=2)
for t in [
    "HIL_TILT_ACTUATOR_CONTROLS 송신: tv_att_control이 발행하는 TiltAngleSetpoint/TiltrotorExtraControls uORB를 구독하여 HIL 모드일 때 동일 주기로 MAVLink 송신",
    "HIL_TILT_STATE 수신: SimulatorMavlink::handle_message_hil_tilt_state()(신규)를 추가하여 수신 필드를 TiltStatus uORB로 발행",
    "HIL 모드에서는 dynamixel_uart 실물 드라이버 퍼블리시를 비활성화하여 TiltStatus 토픽 발행자 중복 방지",
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading("12. uORB ↔ MAVLink 필드 매핑 총괄표", level=1)
doc.add_heading("12.1 채널 A/B (HILS 폐루프)", level=2)
dc.add_table(doc, ["MAVLink 메시지", "방향", "대응 uORB 토픽"], [
    ["HIL_SENSOR", "ENV→FCC", "SensorCombined, sensor_baro, sensor_mag (내부 드라이버 등가 주입)"],
    ["HIL_GPS", "ENV→FCC", "SensorGps"],
    ["HIL_ACTUATOR_CONTROLS", "FCC→ENV", "ActuatorMotors.control[0..3]"],
    ["HIL_TILT_ACTUATOR_CONTROLS (신규)", "FCC→ENV", "TiltAngleSetpoint, TiltrotorExtraControls"],
    ["HIL_TILT_STATE (신규)", "ENV→FCC", "TiltStatus"],
], col_widths_cm=[5.5, 2.5, 8.0])

doc.add_heading("12.2 채널 C (시각화 텔레메트리)", level=2)
dc.add_table(doc, ["MAVLink 메시지", "대응 uORB 토픽", "핵심 필드"], [
    ["GLOBAL_POSITION_INT", "VehicleGlobalPosition", "lat, lon, alt"],
    ["LOCAL_POSITION_NED", "VehicleLocalPosition", "x, y, z, vx, vy, vz, heading"],
    ["ATTITUDE / ATTITUDE_QUATERNION", "VehicleAttitude", "q[4] (w,x,y,z)"],
    ["VFR_HUD", "AirspeedValidated, VehicleLocalPosition", "true_airspeed_m_s, heading, vz"],
    ["WIND_COV", "Wind", "windspeed_north, windspeed_east"],
    ["HIL_TILT_STATE (재송출)", "TiltStatus", "angle[4], angular_velocity[4]"],
], col_widths_cm=[5.0, 5.5, 5.5])

doc.add_heading("13. 타이밍 및 성능 요구사항", level=1)
dc.add_table(doc, ["항목", "요구값", "비고"], [
    ["채널 A→FCC 처리 지연", "≤ 10 ms (수신~EKF2 반영)", "EKF2 지연보상 한계 내"],
    ["채널 B FCC→ENV 지연", "≤ 10 ms", "FDM 적분 스텝과 정합 필요"],
    ["채널 A/B 왕복 지연(RTT)", "≤ 20 ms", ""],
    ["채널 C 갱신 지연(FCC→VIS)", "≤ 50 ms", "시각적 체감 지연 최소화 목적"],
    ["지터(jitter) 허용치", "≤ 5 ms (표준편차 기준)", "고정 주기 발행기(hrt) 사용 권장"],
], col_widths_cm=[5.0, 6.0, 5.0])

doc.add_heading("14. 오류 처리", level=1)
dc.add_table(doc, ["오류 상황", "검출 방법", "처리 방안"], [
    ["HEARTBEAT 두절", "1.5초 이상 미수신", "페일세이프 절차 적용, 링크 상태 UI 경고"],
    ["체크섬 오류", "MAVLink CRC 불일치", "해당 패킷 폐기, 시퀀스 갭 카운터 증가"],
    ["시퀀스 갭", "seq 필드 불연속", "유실률 통계 반영, 임계치 초과 시 경고 로그"],
    ["HIL_TILT_STATE 필드 NaN", "필드값 검증", "직전 유효값 유지(hold-last) 또는 comm_critical 설정"],
], col_widths_cm=[4.0, 5.5, 6.5])

doc.add_heading("15. 연동 시퀀스 (개요)", level=1)
for t in [
    "① [초기화] FCC 부팅 → SYS_HITL=1 파라미터 로드 → MAVLink 링크 오픈",
    "② [연결수립] ENV/VIS ↔ FCC 간 HEARTBEAT 상호 교환",
    "③ [HIL 모드 진입] ENV가 HIL_SENSOR/HIL_GPS/HIL_TILT_STATE 주기 송신 시작 → EKF2 정렬",
    "④ [정상 운용] 채널 A/B 폐루프 지속, 채널 C로 VIS에 상태 지속 송출",
    "⑤ [시나리오 종료] Disarm → SYS_HITL=0 복귀",
]:
    doc.add_paragraph(t, style='List Number')

doc.add_heading("16. 부록 A — 관련 uORB 메시지 원본 필드 발췌", level=1)
doc.add_heading("A.1 VehicleGlobalPosition.msg", level=3)
dc.add_table(doc, ["필드", "타입", "단위"], [
    ["lat, lon", "float64", "degrees"],
    ["alt", "float32", "meters (AMSL)"],
    ["eph, epv", "float32", "meters"],
], col_widths_cm=[6.0, 4.0, 6.0])
doc.add_heading("A.2 VehicleLocalPosition.msg (발췌)", level=3)
dc.add_table(doc, ["필드", "타입", "단위"], [
    ["x, y, z", "float32", "meters (NED)"],
    ["vx, vy, vz", "float32", "m/s (NED)"],
    ["heading", "float32", "rad (-PI..PI)"],
    ["dist_bottom", "float32", "meters"],
], col_widths_cm=[6.0, 4.0, 6.0])
doc.add_heading("A.3 TiltStatus.msg (전체)", level=3)
dc.add_table(doc, ["필드", "타입", "단위/비고"], [
    ["goal_angle[4] / angle[4] / angular_velocity[4]", "float32[4]", "목표각/실제각(deg) / 각속도(deg/s)"],
    ["current[4]", "int16[4]", "전류 (mA)"],
    ["voltage[4] / temperature[4]", "float32[4]", "전압(V) / 온도(°C)"],
    ["moving[4]", "uint8[4]", "이동 중 여부"],
    ["comm_warn[4] / comm_critical[4]", "bool[4]", "통신 상태 플래그"],
], col_widths_cm=[7.0, 3.5, 5.5])

# =====================================================================
# PART III — 비행데이터 분석 파라미터 가이드 (신규)
# =====================================================================
doc.add_page_break()
part_divider(doc, "Ⅲ", "비행데이터 분석 파라미터 가이드",
             "비행 후 ULog를 분석할 때 어떤 파라미터를 봐야 하는가")

doc.add_heading("17. 개요", level=1)
doc.add_heading("17.1 목적", level=2)
doc.add_paragraph(
    "PART Ⅰ·Ⅱ에서 정의한 인터페이스 신호(A-xx/B-xx/C-xx)는 실비행/HILS 중에는 MAVLink 메시지로 "
    "오가지만, 비행이 끝난 뒤에는 FCC가 기록한 ULog(.ulg) 파일 안에 uORB 토픽 형태로 고스란히 남는다. "
    "본 Part는 비행데이터 분석 담당자가 '무엇을 봐야 하는지' 헤매지 않도록, ICD에서 정의한 신호와 "
    "ULog 토픽/필드를 연결하고, 목적별로 그룹핑하여 제시하는 것을 목적으로 한다."
)

doc.add_heading("17.2 근거 및 범위", level=2)
doc.add_paragraph(
    "본 Part의 토픽/필드 목록은 임의로 선정한 것이 아니라, 사내에서 이미 사용 중인 비행로그 변환 도구 "
    "ulog2csv.py의 DEFAULT_TOPIC_FIELDS(기본 비행시험 토픽 목록)를 그대로 근거로 한다. 즉, 아래 목록은 "
    "'이 도구로 CSV 변환 시 기본으로 뽑히는 파라미터 = 비행시험에서 통상적으로 확인해야 하는 파라미터'라는 "
    "전제를 따른다."
)
dc.note(doc, "출처: ulog2csv.py DEFAULT_TOPIC_FIELDS (사내 flight_plotter.py GUI 도구가 동일 정의를 재사용). 신규 토픽 추가 시 본 Part도 함께 갱신 필요.")

doc.add_heading("18. 분석 목적별 파라미터 그룹", level=1)
doc.add_paragraph(
    "아래 표는 uORB 토픽을 8개의 분석 목적 그룹으로 묶고, PART Ⅱ의 인터페이스 신호(9.4절 Master Signal "
    "List)와 대응되는 항목은 신호 ID(A-xx/B-xx/C-xx)를 표기했다. 신호 ID가 없는 항목은 FCC 내부 제어 "
    "루프 전용 데이터로, ICD 외부 인터페이스에는 포함되지 않지만 비행 성능 분석에는 핵심적인 값이다."
)

dc.start_landscape_section(doc)

doc.add_heading("[그룹 1] 위치 / 항법 (Navigation)", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["vehicle_local_position", "x,y,z, vx,vy,vz, heading, dist_bottom, eph, epv", "EKF2가 추정한 로컬 위치/속도/헤딩, 지면고도, 위치정밀도", "PART Ⅱ 9.4절 참고\n(EKF2 산출값)", "eph/epv 급증 = 위치추정 신뢰도 저하\ndist_bottom 불안정 = 지형추종 이상"],
    ["vehicle_local_position_setpoint", "x,y,z, vx,vy,vz, yaw, yawspeed", "위치제어 지령치 — 실제값과의 추종오차 확인용", "N/A (내부 제어루프)", "지령-실제 오차 지속 = 위치제어 성능저하"],
    ["vehicle_global_position", "lat, lon, alt, terrain_alt", "전역 위치, 지형고도", "C-01", "alt-terrain_alt 차이로 실제 AGL 확인"],
    ["vehicle_gps_position", "latitude_deg, longitude_deg, altitude_msl_m, vel_*, satellites_used, fix_type, hdop, eph, epv", "GPS 원시 품질 지표", "A-06 ~ A-08", "fix_type<3, satellites_used<6, hdop>2.0 → GPS 품질 저하 구간"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 2] 자세 / 제어 (Attitude & Control)", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["vehicle_attitude", "q[0..3]", "실제 자세(쿼터니언)", "C-02", "급격한 q 변화 = 자세 이상/진동"],
    ["vehicle_attitude_setpoint", "roll_body, pitch_body, yaw_body, thrust_body[0..2]", "자세 지령치 — 실제 자세와의 추종오차 확인", "N/A (내부 제어루프)", "큰 추종오차 지속 = 제어이득 부적절/외란"],
    ["vehicle_rates_setpoint", "roll, pitch, yaw, thrust_body[0..2]", "각속도 지령치", "N/A (내부 제어루프)", "실제 각속도와 큰 편차 지속 시 확인"],
    ["vehicle_angular_velocity", "xyz[0..2]", "실제 각속도 — 노이즈/진동 확인", "N/A (내부 제어루프)", "고주파 성분 급증 = 진동/구조공진 의심"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 3] 틸트로터 특화 (Tilt Mechanism) — 본 기체 핵심 분석 대상", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["tilt_status", "angle[4], goal_angle[4], current[4]", "실제각 vs 목표각 오차(서보 성능), 전류(과부하/스톨 감지)", "A-09 ~ A-11", "|goal_angle-angle| 지속 이탈 = 서보 성능저하/기계적 걸림\ncurrent 급증 = 과부하/스톨"],
    ["tilt_angle_setpoint", "tilt_fl, tilt_fr, tilt_rl, tilt_rr", "개별 서보 목표각 지령", "B-02", "4개 서보 지령값 간 비대칭 확인"],
    ["tilt_average_angle_setpoint", "tilt_avg_setpoint", "평균 틸트 지령(전이 진행도 지표)", "B-03 (집단값 상응)", "지령 변화율 급격 시 전이 충격 의심"],
    ["tiltrotor_extra_controls", "collective_tilt_normalized_setpoint, collective_thrust_normalized_setpoint", "전이(Transition) 진행률, 추력배분", "B-03", "전이 중 정규값 정체/역행 = 전이 실패 징후"],
    ["vtol_vehicle_status", "vehicle_vtol_state", "MC/FW/TRANSITION 모드 상태", "N/A (내부 상태기)", "전이 전후 구간 분리 기준점으로 사용"],
    ["vehicle_status", "in_transition_mode, in_transition_to_fw, vehicle_type", "전이구간 식별", "N/A (내부 상태기)", "전이 소요시간 비정상 확대 확인"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 4] 추력 / 액추에이터 출력", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["actuator_motors", "control[0..7]", "모터 정규화 출력 명령", "B-01", "특정 모터만 지속 포화(±1) = 고장/편심 의심"],
    ["actuator_servos", "control[0..7]", "서보(조종면 등) 출력 명령", "N/A (내부 제어루프)", "이상 출력/발진 확인"],
    ["actuator_outputs", "output[0..7]", "실제 PWM/DShot 출력값", "N/A (내부 제어루프)", "명령 대비 실제 출력 불일치 = 믹서/ESC 이상"],
    ["control_surface_cmd", "left_aileron, right_aileron, left_ruddervator, right_ruddervator", "고정익 모드 조종면 명령", "N/A (내부 제어루프)", "좌우 비대칭 지속 = 트림/기계 이상"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 5] 대기자료 / 성능 (Air Data & Performance)", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["airspeed_validated", "indicated/calibrated/true_airspeed_m_s", "대기속도 — 실속/구조하중 관련", "C-03", "전이구간 중 최소요구속도(TBD) 미만 = 실속 위험"],
    ["tecs_status", "altitude_sp/reference, height_rate_setpoint/height_rate, true_airspeed_sp/filtered, throttle_sp, pitch_sp_rad", "고정익 고도/속도 제어루프 추종성능", "N/A (내부 제어루프)", "지령-실제 지속 괴리 = 성능저하/바람외란"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 6] 전력 (Power)", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["battery_status", "voltage_v, current_a, remaining, discharged_mah, temperature", "배터리 상태 — 저전압/과전류/저온 경고", "N/A (내부 전력계통)", "voltage_v 급강하, remaining<20%(예시), 온도 이상 = 조기착륙 판단 근거"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 7] 조종 입력 / 비행모드 (RC & Mode)", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["manual_control_setpoint", "roll, pitch, yaw, throttle, flaps", "수동조종 입력값", "D-03", "조종입력과 실제 반응 시차 확인"],
    ["input_rc", "rssi, values[0..7]", "RC 링크 품질 및 원시 채널값", "D-01, D-02", "rssi 저하 = 조종신호 열화, 페일세이프 유발 가능"],
    ["vehicle_status", "arming_state, nav_state, vehicle_type", "암상태/비행모드 전이 시점", "N/A (내부 상태기)", "nav_state 급변 = 페일세이프/모드전환 이벤트"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

doc.add_heading("[그룹 8] 착륙 판정 / 시스템 자원", level=3)
dc.add_table(doc, ["uORB 토픽", "대표 필드", "의미 / 분석 목적", "관련 ICD 신호ID", "점검 포인트 (이상징후)"], [
    ["vehicle_land_detected", "landed, in_ground_effect", "착지/지면효과 판정", "N/A (내부 상태기)", "착지 판정 지연/오탐 = 착륙로직 재검토 필요"],
    ["cpuload", "load, ram_usage", "FCC 연산 부하", "N/A (시스템 자원)", "load 지속 고부하(예시 90%↑) = 제어주기 지연/로그 드롭 가능성"],
], col_widths_cm=[3.6, 6.2, 6.4, 2.8, 6.5])

dc.end_landscape_section(doc)

doc.add_heading("19. 이상징후 점검 체크리스트", level=1)
doc.add_paragraph("비행데이터 분석 시작 시 아래 항목을 우선 스캔하여 이상구간을 빠르게 선별하는 것을 권장한다.")
dc.start_landscape_section(doc)
dc.add_table(doc, ["No", "점검 항목", "관련 파라미터 (토픽.필드)", "판단 기준 (예시, TBD)", "조치 / 해석"], [
    ["1", "GPS 품질 저하", "vehicle_gps_position.fix_type / satellites_used / hdop", "fix_type<3 또는 satellites_used<6 또는 hdop>2.0", "해당 구간은 위치신뢰도 낮음 — 항법 이상 분석 시 우선 배제 또는 별도 표기"],
    ["2", "틸트 서보 추종 오차", "tilt_status.goal_angle[] vs angle[]", "오차가 지속적으로 5° 이상(예시)", "서보 성능저하/기계적 걸림 의심 → 하드웨어 점검"],
    ["3", "틸트 서보 과전류", "tilt_status.current[]", "정격전류의 80%(예시) 이상 지속", "과부하/스톨 가능성 → 즉시 원인 분석"],
    ["4", "전이구간 실속 위험", "airspeed_validated.calibrated_airspeed_m_s + vtol_vehicle_status", "전이 중 CAS가 최소요구속도(TBD) 미만", "실속/추력배분 이상 점검"],
    ["5", "자세 진동/발산", "vehicle_angular_velocity.xyz[], vehicle_attitude.q[]", "고주파 성분 급증, 발산 경향", "제어이득 또는 구조 공진 점검"],
    ["6", "전력 이상", "battery_status.voltage_v / remaining / temperature", "전압 급강하, remaining<20%(예시), 온도 이상", "조기 착륙 판단 근거로 활용"],
    ["7", "RC/링크 열화", "input_rc.rssi + vehicle_status.nav_state", "rssi 저하와 동시에 nav_state 변화(페일세이프 진입)", "링크 이상과 안전모드 전환의 상관관계 확인"],
    ["8", "FCC 연산 부하", "cpuload.load", "부하 90%(예시) 이상 지속", "제어주기 지연/로그 드롭 가능성 — 실시간성 저하 의심"],
    ["9", "모터 포화/불균형", "actuator_motors.control[]", "특정 채널만 지속적으로 ±1 포화", "모터/프롭 손상 또는 무게중심 편심 의심"],
    ["10", "착지 판정 이상", "vehicle_land_detected.landed", "착지 판정 지연 또는 오탐", "착륙로직/지면효과 파라미터 재검토"],
], col_widths_cm=[1.2, 4.5, 6.5, 6.5, 7.3])
dc.note(doc, "판단 기준 수치는 모두 예시(TBD)이며, 실제 기체 성능요구도(ORD)·서보/센서 데이터시트 확정 후 갱신이 필요하다.")
dc.end_landscape_section(doc)

doc.add_heading("20. 로그 추출 방법 (사내 도구 연계)", level=1)
doc.add_paragraph(
    "본 Part의 파라미터 그룹은 사내 ulog2csv.py 스크립트의 기본 추출 대상과 동일하다. 별도 옵션 없이 "
    "실행하면 18장에서 정리한 8개 그룹의 토픽이 그대로 CSV로 변환된다."
)
p = doc.add_paragraph()
run = p.add_run(
    "# 기본 토픽 세트로 변환 (18장 파라미터 그룹과 동일)\n"
    "python ulog2csv.py flight.ulg -o flight.csv\n\n"
    "# 특정 토픽만 추출 (예: 틸트 관련 이상징후 분석)\n"
    "python ulog2csv.py flight.ulg -o tilt_check.csv --messages tilt_status,tiltrotor_extra_controls,vtol_vehicle_status --rate 50\n\n"
    "# 로그에 존재하는 모든 토픽/필드 추출 (파일 용량 큼 — 정밀 조사 시에만 사용)\n"
    "python ulog2csv.py flight.ulg -o flight_all.csv --messages all"
)
run.font.name = "Consolas"; run.font.size = Pt(9.5)
doc.add_paragraph(
    "변환된 CSV는 flight_plotter.py(사내 GUI 도구)로 열어 토픽명 기준으로 그룹핑된 트리에서 원하는 "
    "파라미터를 골라 그래프에 배치할 수 있다."
)

doc.add_heading("21. 부록 B — 상황별 우선 확인 파라미터 (Quick Reference)", level=1)
dc.add_table(doc, ["분석 상황", "우선 확인 파라미터"], [
    ["비행 중 이상 진동/떨림이 보고됨", "vehicle_angular_velocity, tilt_status.current[]"],
    ["전이(천이) 구간에서 불안정 거동", "vtol_vehicle_status, tiltrotor_extra_controls, airspeed_validated, tilt_status"],
    ["항법 이상 / 위치가 튐", "vehicle_gps_position(fix_type, hdop), vehicle_local_position(eph, epv)"],
    ["목표 비행시간 미달", "battery_status"],
    ["조종 반응 지연/끊김", "input_rc.rssi, vehicle_status.nav_state"],
    ["착륙 시 충격/불시착", "vehicle_land_detected, vehicle_local_position.vz, actuator_motors"],
    ["특정 모터/서보 고장 의심", "actuator_motors.control[], tilt_status.current[]/angle[]"],
], col_widths_cm=[6.0, 10.0])
doc.add_paragraph(
    "본 부록은 완결된 진단 도구가 아니라 분석 시작점을 좁히기 위한 참고표이다. 실제 원인 규명은 관련 "
    "파라미터를 시계열로 함께 겹쳐 보고, 필요 시 PART Ⅰ·Ⅱ의 인터페이스 정의(정상 범위/분해능)와 대조하여 "
    "판단한다."
)

doc.save("ICD-PXTR-HILS-000-통합본.docx")
print("Master document saved")
