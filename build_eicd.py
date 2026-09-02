# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import doc_common as dc

doc = Document()
dc.set_base_style(doc)
dc.add_page_number_footer(doc)
dc.enable_auto_update_fields(doc)

dc.add_title_page(
    doc,
    "PX4 기반 틸트로터 VTOL 비행제어시스템",
    "HILS/시각화 연동 외부 인터페이스 통제 문서 (EICD)",
    doc_no="EICD-PXTR-HILS-001",
    rev="A (초안)",
    date_str="2026-08-31",
)

dc.add_revision_history(doc, [
    ["A", "2026-08-31", "최초 발행 (초안) — HILS/시각화 연동 범위 정의", "비행제어SW팀"],
])

doc.add_heading("목차 (Table of Contents)", level=1)
dc.add_toc(doc)
doc.add_page_break()

# =====================================================================
doc.add_heading("1. 개요", level=1)

doc.add_heading("1.1 목적", level=2)
doc.add_paragraph(
    "본 문서는 Airbility 커스텀 PX4 비행제어소프트웨어(OFP: Operational Flight Program)가 "
    "HILS(Hardware-In-the-Loop Simulation) 환경 시뮬레이터 및 시각화(Visualization) 시스템과 "
    "연동하기 위해 필요한 외부 인터페이스(물리적/논리적 연결, 프로토콜, 방향, 주기)를 정의하는 것을 "
    "목적으로 한다. 각 인터페이스에서 교환되는 메시지/필드 단위의 상세 정의는 별도 문서인 "
    "「ICD-PXTR-HILS-001 인터페이스 통제 문서」를 따른다."
)

doc.add_heading("1.2 적용범위", level=2)
doc.add_paragraph(
    "본 문서는 아래 3개 서브시스템 간 외부 인터페이스에 적용된다."
)
for t in [
    "① 비행제어시스템(FCC, Flight Control Computer) — PX4 커스텀 펌웨어 (airbility-dev/px4-custom-firmware, develop/main-v1.15.4 기준)",
    "② HILS 환경 시뮬레이터(ENV) — 비행 동역학(FDM) 및 센서모델을 계산하는 외부 소프트웨어/장비",
    "③ 시각화 시스템(VIS) — 조종사/관제사에게 3D 외부 시계(Out-The-Window) 또는 지도/HUD 화면을 제공하는 소프트웨어",
]:
    doc.add_paragraph(t, style='List Bullet')
doc.add_paragraph(
    "실비행(운용) 모드에서 FCC와 실제 센서/액추에이터(IMU, GPS, 틸트 서보 등) 간 인터페이스는 "
    "본 문서 3.3절(참고용)로 다루되, 상세 하드웨어 ICD는 별도 문서로 관리한다."
)

doc.add_heading("1.3 용어 및 약어 정의", level=2)
dc.add_table(doc, ["약어", "정의"], [
    ["FCC", "Flight Control Computer — 비행제어컴퓨터 (본 프로젝트에서는 PX4 OFP가 탑재되는 보드)"],
    ["OFP", "Operational Flight Program — 비행제어소프트웨어(펌웨어)"],
    ["HILS / HITL", "Hardware-In-the-Loop Simulation — 실제 FCC 하드웨어에 가상 센서 입력을 주입하여 폐루프로 검증하는 시뮬레이션 기법"],
    ["SITL", "Software-In-the-Loop — FCC 소프트웨어를 PC 상에서 실행하는 순수 소프트웨어 시뮬레이션"],
    ["ENV", "HILS 환경 시뮬레이터 (비행 동역학 모델 + 센서 모델 + 지형/기상 모델)"],
    ["VIS", "시각화 시스템 (3D 뷰, 지도, HUD 등)"],
    ["EICD", "External Interface Control Document — 외부 인터페이스 통제 문서(본 문서)"],
    ["ICD", "Interface Control Document — 메시지/필드 단위 상세를 정의하는 인터페이스 통제 문서"],
    ["MAVLink", "PX4/ArduPilot 계열에서 표준으로 사용하는 경량 UAV 통신 프로토콜"],
    ["uORB", "PX4 내부 프로세스 간 발행-구독(pub-sub) 미들웨어 및 그 메시지(토픽)"],
    ["EKF2", "PX4의 확장 칼만필터 기반 상태추정 모듈 (위치/속도/자세 추정)"],
    ["NED", "North-East-Down — 로컬 지평좌표계"],
    ["FRD", "Forward-Right-Down — 기체(Body) 좌표계"],
    ["WGS84", "World Geodetic System 1984 — GPS 위경도 기준 좌표계"],
    ["HIL_*", "MAVLink 표준 HIL(Hardware-In-the-Loop) 메시지 군 (HIL_SENSOR, HIL_GPS 등)"],
], col_widths_cm=[3.0, 13.0])

doc.add_heading("1.4 참고 문서", level=2)
for t in [
    "PX4 사용자/개발자 가이드 (https://docs.px4.io/main/en/)",
    "MAVLink Common/Development Message Set (https://mavlink.io/en/messages/)",
    "airbility-dev/px4-custom-firmware 소스코드 (msg/, src/modules/simulation/, src/modules/mavlink/)",
    "ICD-PXTR-HILS-001 인터페이스 통제 문서 (본 EICD의 하위 문서)",
]:
    doc.add_paragraph(t, style='List Bullet')

# =====================================================================
doc.add_heading("2. 시스템 개요", level=1)

doc.add_heading("2.1 시스템 구성", level=2)
doc.add_paragraph(
    "본 시스템은 실물 FCC 보드(PX4 커스텀 OFP 탑재)를 중심으로, HILS 환경 시뮬레이터가 비행 동역학과 "
    "센서 신호를 계산하여 FCC에 주입하고, FCC는 이를 실제 비행 상황과 동일하게 인식하여 제어연산을 수행한 뒤 "
    "액추에이터(모터/틸트 서보) 명령을 다시 환경 시뮬레이터로 되돌려 폐루프를 형성한다. 동시에 FCC가 "
    "추정한 상태값(위치/속도/자세/틸트각 등)은 시각화 시스템으로 전달되어 외부 환경(시점, 지형, HUD 등)을 "
    "실시간으로 갱신한다."
)
dc.note(doc, "구성도는 별첨 그림 1을 참조. 본 문서에서는 텍스트 블록도로 대체 표기함.")

p = doc.add_paragraph()
run = p.add_run(
    "[FCC: PX4 OFP]\n"
    "   ↑ᴬ HIL 센서/GPS/틸트 피드백 주입 (ENV→FCC)\n"
    "   ↓ᴮ 액추에이터/틸트 명령 출력 (FCC→ENV)\n"
    "[ENV: HILS 환경 시뮬레이터 — FDM/센서모델/지형·기상모델]\n"
    "\n"
    "[FCC: PX4 OFP]\n"
    "   ↓ᶜ 상태 텔레메트리 (FCC→VIS)\n"
    "[VIS: 시각화 시스템 — 3D 뷰/지도/HUD]"
)
run.font.name = "Consolas"; run.font.size = Pt(9.5)

doc.add_heading("2.2 인터페이스 목록", level=2)
dc.add_table(doc, ["ID", "명칭", "연결 시스템", "매체", "방향", "용도"], [
    ["EICD-01", "FCC ↔ HILS 환경 시뮬레이터\n(센서 주입 / 액추에이터 피드백)", "FCC ↔ ENV", "Ethernet(UDP) 또는 USB 시리얼", "양방향", "HILS 폐루프 핵심 인터페이스"],
    ["EICD-02", "FCC → 시각화 시스템\n(상태 텔레메트리)", "FCC → VIS", "Ethernet(UDP/TCP) 또는 시리얼", "단방향\n(FCC→VIS)", "외부 환경/HUD 갱신용 상태 송출"],
    ["EICD-03", "FCC ↔ 틸트 액추에이터 하드웨어\n(Dynamixel, UART)", "FCC ↔ 틸트 서보 4식", "UART (반이중, RS-485/TTL)", "양방향", "실비행 시 실제 틸트 구동/피드백\n(HILS 모드에서는 EICD-01의\nHIL_TILT_* 메시지로 대체됨)"],
    ["EICD-04\n(참고)", "ENV ↔ VIS\n(환경/지형 상태 공유)", "ENV ↔ VIS", "구현 방식에 따름", "양방향", "본 EICD 범위 밖(ENV·VIS 내부 구현),\n필요 시 별도 ICD로 관리"],
    ["EICD-05", "조종기(RC/조이스틱) → FCC\n(검증용 입력)", "RC Tx/GCS → FCC", "RC 수신기 UART\n또는 MAVLink(GCS 경유)", "단방향\n(조종기→FCC)", "HILS 중 조종 반응 검증용"],
], col_widths_cm=[1.8, 4.2, 3.0, 3.5, 2.0, 3.5])

# =====================================================================
doc.add_heading("3. 인터페이스별 상세 정의", level=1)

doc.add_heading("3.1 EICD-01: FCC ↔ HILS 환경 시뮬레이터", level=2)

doc.add_heading("3.1.1 물리 인터페이스", level=3)
dc.add_table(doc, ["항목", "내용"], [
    ["연결 매체", "1순위: Ethernet (UDP/IP) — 개발/실험실 환경 권장\n2순위: USB-시리얼 (FCC의 TELEM 포트) — 실기체 텔레메트리 포트를 그대로 사용할 경우"],
    ["네트워크 포트(UDP 예시)", "FCC 수신(센서 주입): UDP 4560\nFCC 송신(액추에이터 피드백): UDP 14580\n(PX4 SITL/HITL 기본 포트 관례를 계승, 실제 값은 사내 네트워크 설계서에서 최종 확정)"],
    ["시리얼 파라미터(대안)", "Baudrate 921600 bps, 8N1, HW 흐름제어 없음"],
    ["케이블/커넥터", "Ethernet: RJ45 / USB-시리얼: Micro-USB 또는 JST-GH (보드 사양에 따름)"],
], col_widths_cm=[4.0, 12.0])

doc.add_heading("3.1.2 논리 인터페이스 (프로토콜)", level=3)
doc.add_paragraph(
    "본 인터페이스는 PX4가 표준으로 지원하는 MAVLink v2 HIL 메시지 세트를 기반으로 하며, "
    "PX4 소스 내 simulator_mavlink 모듈(src/modules/simulation/simulator_mavlink/SimulatorMavlink.cpp)에서 "
    "그대로 처리 가능한 방식을 채택한다. 이는 별도의 통신 스택을 새로 개발할 필요가 없고, 기존 SITL "
    "타깃(jMAVSim, Gazebo 등)과 동일한 인터페이스로 검증이 가능하다는 장점이 있다."
)
dc.add_table(doc, ["방향", "논리 채널", "대표 메시지", "설명"], [
    ["ENV → FCC", "센서 주입", "HIL_SENSOR, HIL_GPS,\n(옵션) HIL_OPTICAL_FLOW,\nDISTANCE_SENSOR", "IMU/GPS/기압/지자기 등 가상 센서값을 FCC의 EKF2에 주입"],
    ["ENV → FCC", "틸트 피드백 (커스텀)", "HIL_TILT_STATE (신규)", "4개 틸트 서보의 실제 각도/각속도/전류/온도 피드백을 시뮬레이션 값으로 대체 주입"],
    ["FCC → ENV", "액추에이터 출력", "HIL_ACTUATOR_CONTROLS", "모터 추력 명령(4식) — 환경 시뮬레이터의 FDM 입력으로 사용"],
    ["FCC → ENV", "틸트 명령 (커스텀)", "HIL_TILT_ACTUATOR_CONTROLS (신규)", "4개 틸트 서보 목표각(setpoint) 명령 — 환경 시뮬레이터가 서보 동特성 모델로 응답 계산"],
    ["FCC → ENV", "상태 비교/로깅(옵션)", "HIL_STATE_QUATERNION", "FCC 자체 추정치와 진리값(ground truth) 비교용, 제어 루프에는 필수 아님"],
], col_widths_cm=[2.2, 2.8, 4.0, 7.0])
dc.note(doc, "HIL_TILT_STATE, HIL_TILT_ACTUATOR_CONTROLS는 MAVLink 표준셋에 없는 신규 메시지로, 상세 정의 및 임시 메시지 ID는 ICD 문서 4장을 따른다. 배포 전 사내/외부 시스템과 ID 충돌 여부 확인 필수.")

doc.add_heading("3.1.3 데이터 갱신 주기 요구사항", level=3)
dc.add_table(doc, ["데이터", "최소 주기", "권장 주기", "비고"], [
    ["IMU (가속도/각속도)", "150 Hz", "250 Hz", "EKF2 예측 스텝과 정합되어야 함"],
    ["GPS", "5 Hz", "10 Hz", "실제 GPS 모듈 갱신율과 동일하게 설정"],
    ["기압/자기장", "10 Hz", "50 Hz", ""],
    ["틸트 피드백 (HIL_TILT_STATE)", "50 Hz", "100 Hz", "Dynamixel 실제 피드백 주기(TiltStatus)와 정합"],
    ["액추에이터 출력 (FCC→ENV)", "FCC 제어루프 주기와 동일", "250~400 Hz", "mc/tv 제어 루프 주기에 종속"],
], col_widths_cm=[5.0, 3.0, 3.0, 5.0])

doc.add_heading("3.1.4 동기화 및 시간 기준", level=3)
doc.add_paragraph(
    "모든 메시지는 마이크로초(usec) 단위의 단조증가 타임스탬프를 포함한다. FCC와 ENV는 세션 시작 시 "
    "하트비트(HEARTBEAT) 교환을 통해 시각 기준(epoch)을 정렬하며, 이후에는 각자의 로컬 클록을 기준으로 "
    "타임스탬프를 발행한다. 왕복 지연(RTT)은 20 ms 이내를 목표로 하며, 이를 초과할 경우 EKF2 지연 보상 "
    "한계를 벗어나 상태추정 발산 위험이 있으므로 3.1.5절 페일세이프를 따른다."
)

doc.add_heading("3.2 EICD-02: FCC → 시각화 시스템", level=2)
dc.add_table(doc, ["항목", "내용"], [
    ["물리 인터페이스", "Ethernet(UDP/TCP), 필요 시 시리얼 텔레메트리 포트 분기"],
    ["논리 프로토콜", "MAVLink v2 텔레메트리 스트림 (GLOBAL_POSITION_INT, LOCAL_POSITION_NED, ATTITUDE, VFR_HUD, WIND_COV) + 커스텀 HIL_TILT_STATE 재송출"],
    ["방향", "단방향 (FCC → VIS). 다만 VIS에서 뷰포인트 전환 등 명령이 필요할 경우 향후 MAV_CMD 기반 채널 추가 검토"],
    ["권장 갱신 주기", "위치/속도: 20~50 Hz, 자세: 50~100 Hz (시각화 엔진 측 보간을 전제로 함), 틸트 상태: 20 Hz"],
    ["연결 방식", "FCC의 두 번째 MAVLink 인스턴스(-i 1) 또는 UDP 브로드캐스트로 SITL 인터페이스(EICD-01)와 물리적으로 분리 권장"],
], col_widths_cm=[4.0, 12.0])

doc.add_heading("3.3 EICD-03: FCC ↔ 틸트 액추에이터 하드웨어 (참고, 실비행 기준)", level=2)
doc.add_paragraph(
    "본 절은 실비행(실기체) 구성에서의 물리 인터페이스를 참고용으로 기술한다. HILS 모드에서는 본 채널이 "
    "비활성화되고, 동일한 역할을 EICD-01의 HIL_TILT_ACTUATOR_CONTROLS / HIL_TILT_STATE 메시지가 대신한다."
)
dc.add_table(doc, ["항목", "내용"], [
    ["매체", "UART (dynamixel_uart, uart_reader 모듈에서 처리)"],
    ["대상 장치", "틸트 서보 4식 (FL/FR/RL/RR, Dynamixel 프로토콜)"],
    ["FCC → 서보", "목표각 명령 (TiltAngleSetpoint uORB → Dynamixel Goal Position)"],
    ["서보 → FCC", "각도/각속도/전류/온도/통신상태 피드백 (TiltStatus uORB로 발행)"],
    ["HILS 대체 매핑", "TiltAngleSetpoint → HIL_TILT_ACTUATOR_CONTROLS (송신)\nHIL_TILT_STATE(수신) → TiltStatus uORB로 발행하여 실서보와 동일한 방식으로 상위 제어 모듈에 공급"],
], col_widths_cm=[4.0, 12.0])

doc.add_heading("3.4 EICD-05: 조종기(RC/조이스틱) → FCC (검증용 입력)", level=2)
doc.add_paragraph(
    "HILS 시나리오 중 FCC가 실제 조종 입력에 정상적으로 반응하는지 검증하기 위해, 실제 RC 송신기 또는 "
    "지상시스템(GCS)의 조이스틱 입력을 MAVLink로 FCC에 전달하는 별도 검증 채널이다. 이 채널은 비행 "
    "동역학(ENV)이나 시각화(VIS)와는 무관하며, 오직 '조종 입력 → FCC 응답'의 정합성 확인이 목적이다."
)
dc.add_table(doc, ["항목", "내용"], [
    ["물리 인터페이스", "① RC 수신기를 FCC에 직결(UART/SBUS 등, 실비행과 동일 경로) 또는\n② GCS/조이스틱 PC에서 MAVLink로 송신(EICD-01/02와 동일 네트워크 경유)"],
    ["논리 프로토콜", "MAVLink v2 — RC_CHANNELS(#65, 원시 채널값) 또는 MANUAL_CONTROL(#69, 정규화 입력, 조이스틱 대안)"],
    ["방향", "단방향 (조종기/GCS → FCC)"],
    ["용도", "FCC의 manual_control_setpoint, input_rc uORB 반응을 실제 조종 입력과 대조 검증"],
    ["권장 갱신 주기", "RC 수신기 자체 샘플링 50 Hz(예시) / MAVLink 전송 10 Hz(예시, QGC 기본 스트림레이트 기준)"],
], col_widths_cm=[4.0, 12.0])

# =====================================================================
doc.add_heading("4. 좌표계, 단위 및 시간 기준", level=1)
dc.add_table(doc, ["항목", "정의"], [
    ["위치(전역)", "WGS84 위도/경도(도, deg), 고도는 AMSL 기준(m)"],
    ["위치(로컬)", "NED(North-East-Down) 좌표계, FCC 부팅/EKF2 초기화 시점을 원점으로 함 (단위 m)"],
    ["자세", "쿼터니언 q=(w,x,y,z), Hamilton 컨벤션, FRD(기체) → NED(지면) 회전으로 정의"],
    ["각속도", "FRD 기체좌표계 기준 rad/s"],
    ["속도", "NED 좌표계 기준 m/s"],
    ["틸트각", "서보 축 기준 deg (0°: 수직/멀티콥터 모드, 90°: 수평/고정익 모드 — 부호/기준은 ICD 4.2절 표 준수)"],
    ["시간", "메시지 내 time_usec: 시스템 부팅 후 경과 마이크로초(단조증가). UTC 절대시각이 필요한 경우 GPS 메시지의 time_utc_usec 별도 사용"],
], col_widths_cm=[3.5, 12.5])
dc.note(doc, "MAVLink ATTITUDE 등 일부 표준 메시지는 각도를 rad 단위로 정의하므로, deg 기준의 커스텀 틸트 메시지와 혼동하지 않도록 ICD에 단위를 필드마다 명시함.")

# =====================================================================
doc.add_heading("5. 운용 모드 및 천이", level=1)
doc.add_heading("5.1 모드 개요", level=2)
dc.add_table(doc, ["모드", "설명", "센서/액추에이터 경로"], [
    ["실비행 (FLIGHT)", "실제 기체에서 운용", "실 IMU/GPS/틸트 서보 (EICD-03 경로)"],
    ["SITL", "FCC 소프트웨어를 PC에서 실행, 순수 소프트웨어 시뮬레이션", "PX4 내장 시뮬레이터(jMAVSim/Gazebo 등)"],
    ["HITL/HILS (본 문서 대상)", "실물 FCC 보드 + 외부 환경 시뮬레이터", "EICD-01 (HIL_* 메시지 경로), 틸트는 HIL_TILT_* 로 대체"],
], col_widths_cm=[3.5, 6.5, 6.0])

doc.add_heading("5.2 모드 진입 절차 (개요)", level=2)
for t in [
    "① FCC 파라미터 SYS_HITL = 1 설정 (HITL 모드 활성화)",
    "② FCC ↔ ENV 간 MAVLink 링크 연결 및 HEARTBEAT 교환 확인",
    "③ ENV로부터 HIL_SENSOR/HIL_GPS 수신 시작 확인 (FCC 로그 상 EKF2 정렬 여부 확인)",
    "④ 틸트 인터페이스 정상 여부 확인 (HIL_TILT_STATE 수신 → TiltStatus uORB 값 확인)",
    "⑤ 안전 절차에 따라 Arming 후 HILS 시나리오 시작",
]:
    doc.add_paragraph(t, style='List Number')

doc.add_heading("5.3 페일세이프", level=2)
dc.add_table(doc, ["이상 상황", "판단 기준", "FCC 동작"], [
    ["ENV 링크 두절", "HEARTBEAT 미수신 1.5초 이상", "HITL 데이터 소스 무효 처리, 콘솔/로그 경고, (설정 시) 강제 Disarm"],
    ["센서 값 미갱신", "HIL_SENSOR 미수신 500ms 이상", "EKF2 타임아웃 처리, 안전상태 전이"],
    ["틸트 피드백 두절", "HIL_TILT_STATE 미수신 300ms 이상", "TiltStatus comm_loss 카운터 증가 → comm_warn/comm_critical 플래그 설정, tv_att_control 안전모드 전이 검토"],
], col_widths_cm=[4.0, 5.5, 6.5])

# =====================================================================
doc.add_heading("6. 인터페이스 검증 및 형상관리", level=1)
doc.add_heading("6.1 검증 방법", level=2)
for t in [
    "루프백 테스트: ENV 대신 더미 송신기로 HIL_SENSOR/HIL_GPS를 주입하여 FCC 로그(EKF2 status)로 정상 소비 여부 확인",
    "필드 단위 검증: 각 메시지 필드값을 uORB 리스너(listener 커맨드)로 대조하여 스케일/단위/부호 일치 확인",
    "폐루프 검증: 간단한 계단입력(step input)에 대해 ENV의 FDM 응답과 FCC 제어출력이 상호 발산하지 않는지 확인",
    "틸트 인터페이스 검증: HIL_TILT_STATE에 알려진 각도를 주입하고 TiltStatus 및 상위 제어(tv_att_control) 반응 확인",
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading("6.2 개정 및 형상관리", level=2)
doc.add_paragraph(
    "본 문서는 비행제어SW팀이 형상관리 주체이며, 인터페이스 변경(신규 메시지 추가, 포트/주기 변경 등)이 "
    "발생할 경우 ICD 문서와 동시 개정하고 개정이력에 반영한다. 변경 시 관련 부서(환경 시뮬레이터 담당, "
    "시각화 담당)의 리뷰 및 승인을 거쳐야 한다."
)

# =====================================================================
doc.add_heading("7. 부록: 관련 소스 모듈 매핑", level=1)
dc.add_table(doc, ["인터페이스", "관련 PX4 소스 모듈/경로"], [
    ["EICD-01 (HIL 센서/액추에이터)", "src/modules/simulation/simulator_mavlink/SimulatorMavlink.cpp (HIL_SENSOR, HIL_GPS, HIL_STATE_QUATERNION 처리 / HIL_ACTUATOR_CONTROLS 송신)"],
    ["EICD-01 (개별 센서 시뮬레이션 참고)", "src/modules/simulation/sensor_gps_sim, sensor_baro_sim, sensor_mag_sim, sensor_airspeed_sim, battery_simulator"],
    ["EICD-02 (텔레메트리)", "src/modules/mavlink/streams/*.hpp (GLOBAL_POSITION_INT, LOCAL_POSITION_NED, ATTITUDE, VFR_HUD 등)"],
    ["EICD-03 (실 틸트 서보)", "src/modules/dynamixel_uart/, src/modules/uart_reader/, msg/TiltStatus.msg, msg/TiltAngleSetpoint.msg"],
    ["상위 제어(틸트로터)", "src/modules/tv_att_control/, src/modules/tv_pos_control/, src/modules/tv_control_allocator/"],
    ["Gazebo 연동 참고", "src/modules/simulation/gz_bridge (물리엔진 브리지 설계 참고용)"],
], col_widths_cm=[4.5, 11.5])

doc.save("EICD-PXTR-HILS-001.docx")
print("EICD saved")
