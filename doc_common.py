# -*- coding: utf-8 -*-
"""Shared helpers for building ICD/EICD Word documents with python-docx."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

NAVY = RGBColor(0x1F, 0x33, 0x64)
HEADER_FILL = "1F3364"
ALT_FILL = "EEF1F7"
FONT_KR = "맑은 고딕"
FONT_EN = "Calibri"


def set_base_style(doc):
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_KR)

    for hname, size, color in [
        ('Heading 1', 16, NAVY), ('Heading 2', 13, NAVY),
        ('Heading 3', 11.5, RGBColor(0, 0, 0)), ('Heading 4', 10.5, RGBColor(0, 0, 0)),
    ]:
        try:
            hs = doc.styles[hname]
        except KeyError:
            continue
        hs.font.name = FONT_EN
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True
        rpr = hs.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_KR)

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9.5, color=None, align=None, font_kr=FONT_KR):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            if align is not None:
                p.alignment = align
        run = p.add_run(line)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = FONT_EN
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_kr)
        if color:
            run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, col_widths_cm=None, header_fill=HEADER_FILL, zebra=True, repeat_header=False):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        set_cell_text(hdr[i], h, bold=True, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    if repeat_header:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), 'true')
        trPr.append(tblHeader)
    for r_idx, row in enumerate(rows):
        table_row = table.add_row()
        trPr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)
        cells = table_row.cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9)
            if zebra and r_idx % 2 == 1:
                shade_cell(cells[i], ALT_FILL)
    if col_widths_cm:
        table.autofit = False
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
        for i, w in enumerate(col_widths_cm):
            table.columns[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    txt = OxmlElement('w:t'); txt.text = "목차를 보려면 Word에서 필드 업데이트(F9)를 실행하세요."
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep)
    r2 = OxmlElement('w:r'); r2.append(txt)
    run._r.addnext(r2)
    r3 = OxmlElement('w:r'); r3.append(fld_end)
    r2.addnext(r3)


def add_title_page(doc, title_kr, title_sub, doc_no, rev, date_str, classification="내부 검토용 / DRAFT"):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(classification)
    run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_kr)
    run.font.size = Pt(24); run.font.bold = True; run.font.color.rgb = NAVY
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_KR)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_sub)
    run.font.size = Pt(14); run.font.bold = False; run.font.color.rgb = RGBColor(0x40,0x40,0x40)

    for _ in range(6):
        doc.add_paragraph()

    meta = [
        ("문서번호 (Doc No.)", doc_no),
        ("개정번호 (Revision)", rev),
        ("발행일 (Date)", date_str),
        ("발행처", "Airbility Co., Ltd. — 비행제어 SW팀"),
        ("적용 대상", "PX4 커스텀 펌웨어 (airbility-dev/px4-custom-firmware, develop/main-v1.15.4)"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit = False
    for i, (k, v) in enumerate(meta):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(4.5); c1.width = Cm(10.5)
        shade_cell(c0, HEADER_FILL)
        set_cell_text(c0, k, bold=True, color=RGBColor(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(c1, v)
    doc.add_page_break()


def add_revision_history(doc, rows):
    doc.add_heading("개정 이력 (Revision History)", level=1)
    add_table(doc, ["Rev", "일자", "변경 내용", "작성자"], rows, col_widths_cm=[1.5, 2.5, 9.5, 2.5])


def start_landscape_section(doc):
    """Start a new section in landscape orientation. Returns the section."""
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width, new_section.page_height = doc.sections[0].page_height, doc.sections[0].page_width
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)
    new_section.top_margin = Cm(1.8)
    new_section.bottom_margin = Cm(1.8)
    return new_section


def end_landscape_section(doc):
    """Return to portrait orientation matching the document's base page size."""
    new_section = doc.add_section()
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.left_margin = Cm(2.2)
    new_section.right_margin = Cm(2.2)
    new_section.top_margin = Cm(2.0)
    new_section.bottom_margin = Cm(2.0)
    return new_section


def enable_auto_update_fields(doc):
    settings = doc.settings.element
    upd = OxmlElement('w:updateFields')
    upd.set(qn('w:val'), 'true')
    settings.append(upd)


PARAM_LIST_HEADERS = ["No", "신호명(데이터)", "구분", "전송경로", "물리적 범위 (Range)",
                       "갱신주기\n(컴퓨팅 / 계산 / 전송)", "데이터 비트폭 / 인코딩",
                       "통신방법 · 엔디언", "단위", "확정상태"]
PARAM_LIST_COL_WIDTHS = [0.9, 3.0, 2.6, 2.6, 3.0, 4.2, 3.4, 3.0, 1.1, 1.7]

MAVLINK_AB_D_COMM = "MAVLink v2 (UDP, EICD-01)\nLittle-Endian\n(MAVLink 표준)"
MAVLINK_C_COMM = "MAVLink v2 (UDP/TCP, EICD-02)\nLittle-Endian\n(MAVLink 표준)"

PARAM_LIST_ROWS = [
    # --- Channel A: ENV -> FCC, HILS 센서/틸트 피드백 주입 ---
    ["A-01", "체축 가속도\n(Xacc,Yacc,Zacc)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_SENSOR)",
     "±16 g", "컴퓨팅: 1000 Hz(예시, ENV FDM 적분)\n계산: 1000 Hz(예시, IMU 모델)\n전송: 250 Hz",
     "float32 (MAVLink)\n※원 IMU 16bit ADC 가정 시\nLSB≈0.49 mg", MAVLINK_AB_D_COMM, "m/s²", "[TBD-센서사양]"],
    ["A-02", "체축 각속도\n(Xgyro,Ygyro,Zgyro)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_SENSOR)",
     "±2000 deg/s", "컴퓨팅: 1000 Hz(예시)\n계산: 1000 Hz(예시)\n전송: 250 Hz",
     "float32\n※16bit 가정 시 LSB≈0.061 deg/s", MAVLINK_AB_D_COMM, "rad/s", "[TBD-센서사양]"],
    ["A-03", "지자기\n(Xmag,Ymag,Zmag)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_SENSOR)",
     "±8 gauss", "컴퓨팅: 100 Hz(예시)\n계산: 100 Hz(예시)\n전송: 50 Hz",
     "float32", MAVLINK_AB_D_COMM, "gauss", "[TBD-센서사양]"],
    ["A-04", "절대기압\n(abs_pressure)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_SENSOR)",
     "300 ~ 1100 hPa\n(해수면~약12,000 m)", "컴퓨팅: 100 Hz(예시)\n계산: 100 Hz(예시)\n전송: 50 Hz",
     "float32\n※원 바로센서 24bit 가정", MAVLINK_AB_D_COMM, "hPa", "[TBD-센서사양]"],
    ["A-05", "대기자료센서 차압\n(diff_pressure, Pt-Ps)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_SENSOR)",
     "20 ~ 50 inHg\n(예시값, ≈68~169 hPa)",
     "컴퓨팅: 100 Hz(예시, ENV FDM)\n계산: 8 Hz(실 ADS 센서 특성,\n사용자 제시값)\n전송: 8 Hz",
     "8 bit (예시, 원 센서 출력)\n→ FCC 수신 시 float32(hPa)\n로 환산", MAVLINK_AB_D_COMM, "inHg\n(내부 hPa)", "[TBD-센서사양]\n(사용자 예시 반영)"],
    ["A-06", "GPS 위도/경도\n(lat, lon)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_GPS)",
     "위도 ±90°, 경도 ±180°", "컴퓨팅: 1000 Hz(예시, ENV FDM)\n계산: 10 Hz(GPS 모듈 특성)\n전송: 10 Hz",
     "int32, degE7\n(분해능 1e-7°≈1.1 cm)", MAVLINK_AB_D_COMM, "deg", "프로토콜 고정"],
    ["A-07", "GPS 고도 (alt)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_GPS)",
     "-500 ~ 6,000 m (AMSL)", "컴퓨팅: 1000 Hz(예시)\n계산: 10 Hz\n전송: 10 Hz",
     "int32, mm (1 mm 분해능)", MAVLINK_AB_D_COMM, "m", "[TBD-성능요구도]"],
    ["A-08", "GPS 속도 (vn, ve, vd)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_GPS)",
     "±60 m/s (예시)", "컴퓨팅: 1000 Hz(예시)\n계산: 10 Hz\n전송: 10 Hz",
     "int16, cm/s (1 cm/s 분해능)", MAVLINK_AB_D_COMM, "m/s", "[TBD-성능요구도]"],
    ["A-09", "틸트 서보 각도 피드백\n(FL/FR/RL/RR)", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_TILT_STATE, 신규)",
     "0 ~ 100 deg\n(기구 틸트 범위, 예시)", "컴퓨팅: 1000 Hz(예시, 서보모델)\n계산: 100 Hz(엔코더 refresh)\n전송: 100 Hz",
     "float32\n※12bit 엔코더 가정 시\nLSB≈0.088°", MAVLINK_AB_D_COMM, "deg", "[TBD-기구사양]"],
    ["A-10", "틸트 서보 각속도", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_TILT_STATE)",
     "±180 deg/s (예시)", "컴퓨팅: 1000 Hz(예시)\n계산: 100 Hz\n전송: 100 Hz",
     "float32", MAVLINK_AB_D_COMM, "deg/s", "[TBD-서보사양]"],
    ["A-11", "틸트 서보 전류", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_TILT_STATE)",
     "0 ~ 3000 mA (예시)", "컴퓨팅: 1000 Hz(예시)\n계산: 50 Hz\n전송: 50 Hz",
     "int16, 1 mA 분해능", MAVLINK_AB_D_COMM, "mA", "[TBD-서보사양]"],
    ["A-12", "틸트 서보 온도", "① HILS 센서 주입\n(ENV→FCC)", "ENV → FCC OFP\n(HIL_TILT_STATE)",
     "-20 ~ 85 °C (예시)", "컴퓨팅: 10 Hz(예시, 열모델)\n계산: 10 Hz\n전송: 10 Hz",
     "float32", MAVLINK_AB_D_COMM, "°C", "[TBD-서보사양]"],

    # --- Channel B: FCC -> ENV, 액추에이터/틸트 명령 출력 ---
    ["B-01", "모터 추력 명령 ×4\n(FL/FR/RL/RR)", "② 액추에이터 출력\n(FCC→ENV)", "FCC OFP → ENV\n(HIL_ACTUATOR_CONTROLS)",
     "-1 ~ 1 (정규화)\nNaN = Disarm", "컴퓨팅: 400 Hz(예시, rate 제어루프)\n계산: 400 Hz(actuator_motors 발행)\n전송: 250~400 Hz",
     "float32", MAVLINK_AB_D_COMM, "무차원", "프로토콜 고정"],
    ["B-02", "틸트 목표각 명령 ×4\n(FL/FR/RL/RR)", "② 액추에이터 출력\n(FCC→ENV)", "FCC OFP → ENV\n(HIL_TILT_ACTUATOR_CONTROLS, 신규)",
     "0 ~ 100 deg (기구 한계, 예시)", "컴퓨팅: 250 Hz(예시, tv_att_control)\n계산: 100 Hz(TiltAngleSetpoint 발행)\n전송: 100 Hz",
     "float32", MAVLINK_AB_D_COMM, "deg", "[TBD-기구사양]"],
    ["B-03", "집단 틸트 정규화 명령\n(collective_tilt_norm)", "② 액추에이터 출력\n(FCC→ENV)", "FCC OFP → ENV\n(HIL_TILT_ACTUATOR_CONTROLS)",
     "0(수직) ~ 1(수평)", "컴퓨팅: 250 Hz(예시)\n계산: 100 Hz\n전송: 100 Hz",
     "float32", MAVLINK_AB_D_COMM, "무차원", "프로토콜 고정"],

    # --- Channel C: FCC -> VIS(SIM PC), 배경(환경) 표시용 텔레메트리 ---
    ["C-01", "전역 위치\n(위도/경도/고도)", "③ SIM PC 배경 표시용\n(FCC→VIS)", "FCC OFP → VIS\n(GLOBAL_POSITION_INT)",
     "위경도 전역, 고도 -500~6,000 m", "컴퓨팅: 200 Hz(예시, EKF2)\n계산: 50 Hz(vehicle_global_position 발행)\n전송: 10~30 Hz",
     "int32 (degE7 / mm)", MAVLINK_C_COMM, "deg, m", "[TBD-성능요구도]"],
    ["C-02", "자세 쿼터니언\n(q0..q3)", "③ SIM PC 배경 표시용\n(FCC→VIS)", "FCC OFP → VIS\n(ATTITUDE_QUATERNION)",
     "단위쿼터니언 (‖q‖=1)", "컴퓨팅: 200~400 Hz(예시, EKF2)\n계산: 200 Hz(vehicle_attitude 발행)\n전송: 50~100 Hz",
     "float32 ×4", MAVLINK_C_COMM, "무차원", "프로토콜 고정"],
    ["C-03", "대기속도 (IAS/TAS)", "③ SIM PC 배경 표시용\n(FCC→VIS)", "FCC OFP → VIS\n(VFR_HUD)",
     "0 ~ 60 m/s (예시)", "컴퓨팅: 100 Hz(예시)\n계산: 100 Hz(airspeed_validated 발행)\n전송: 10~20 Hz",
     "float32", MAVLINK_C_COMM, "m/s", "[TBD-성능요구도]"],
    ["C-04", "방위각 (heading)", "③ SIM PC 배경 표시용\n(FCC→VIS)", "FCC OFP → VIS\n(VFR_HUD / GLOBAL_POSITION_INT)",
     "0 ~ 360 deg", "컴퓨팅: 50 Hz(예시, 파생값)\n계산: 50 Hz\n전송: 10~20 Hz",
     "uint16, cdeg (0.01° 분해능)", MAVLINK_C_COMM, "deg", "프로토콜 고정"],
    ["C-05", "풍향/풍속 (N/E 성분)", "③ SIM PC 배경 표시용\n(FCC→VIS)", "FCC OFP → VIS\n(WIND_COV)",
     "±30 m/s (예시)", "컴퓨팅: 10 Hz(예시, wind estimator)\n계산: 10 Hz\n전송: 1~5 Hz",
     "float32", MAVLINK_C_COMM, "m/s", "[TBD-성능요구도]"],
    ["C-06", "틸트 상태 (재송출)", "③ SIM PC 배경 표시용\n(FCC→VIS)", "FCC OFP → VIS\n(HIL_TILT_STATE 재송출)",
     "0 ~ 100 deg (예시)", "컴퓨팅: A-09와 동일\n계산: A-09와 동일(100 Hz)\n전송: 20 Hz",
     "float32 ×4", MAVLINK_C_COMM, "deg", "[TBD-기구사양]"],

    # --- Channel D (신규): 조종기 -> FCC, 검증용 입력 ---
    ["D-01", "RC 채널 원시값 ×8\n(chan1..8_raw)", "④ 조종기 입력 검증용\n(RC/MAVLink→FCC)", "조종기(RC Tx/조이스틱)\n→ FCC OFP\n(RC_CHANNELS)",
     "1000 ~ 2000 µs\n(PWM 등가값)", "컴퓨팅: 50 Hz(예시, RC 수신기 샘플링)\n계산: 50 Hz(MAVLink 브리지)\n전송: 10 Hz(예시, QGC 기본 스트림)",
     "uint16 × 8채널", MAVLINK_AB_D_COMM, "µs", "[TBD-장비사양]"],
    ["D-02", "RC 링크 품질 (RSSI)", "④ 조종기 입력 검증용\n(RC/MAVLink→FCC)", "조종기 → FCC OFP\n(RC_CHANNELS.rssi)",
     "0 ~ 100 % (또는 0~255 raw)", "컴퓨팅: 50 Hz(예시)\n계산: 50 Hz\n전송: 10 Hz(예시)",
     "uint8", MAVLINK_AB_D_COMM, "%", "[TBD-장비사양]"],
    ["D-03", "조이스틱 정규화 입력\n(x, y, z, r)", "④ 조종기 입력 검증용\n(RC/MAVLink→FCC)", "GCS/조이스틱 → FCC OFP\n(MANUAL_CONTROL, 대안)",
     "-1000 ~ 1000", "컴퓨팅: 50 Hz(예시)\n계산: 50 Hz\n전송: 10~20 Hz(예시)",
     "int16 × 4", MAVLINK_AB_D_COMM, "무차원", "[TBD-장비사양]"],
]


def add_master_parameter_list(doc, heading_text="ICD 파라미터 목록 (통합)", heading_level=2, number_prefix=""):
    """Emit the single unified parameter list table (Channels A/B/C/D) in a landscape section."""
    doc.add_heading(heading_text, level=heading_level)
    doc.add_paragraph(
        "본 절은 이전까지 채널별로 나뉘어 있던 인터페이스 신호 목록을 하나의 표로 통합한 것이다. "
        "'구분' 열은 신호의 용도를 4가지로 분류한다: "
        "① HILS 센서 주입(ENV가 FCC에 가상 센서값을 주입), "
        "② 액추에이터 출력(FCC가 ENV로 모터/틸트 명령을 송신), "
        "③ SIM PC 배경(환경) 표시용(FCC가 시각화 PC의 화면/환경 갱신을 위해 송신하는 상태값), "
        "④ 조종기 입력 검증용(실제/가상 조종기가 MAVLink로 FCC에 보내 조종 반응을 검증하기 위한 값). "
        "모든 신호에는 전송경로, 물리적 범위, 갱신주기(컴퓨팅/계산/전송 3단계), 데이터 비트폭/인코딩, "
        "통신방법과 엔디언까지 명시하여 별도 설명 없이도 구현 가능하도록 한다. "
        "예) '대기자료센서(Pt)가 ENV에서 FCC OFP로 계산 8 Hz·전송 8 Hz로 20~50 inHg 범위의 값을 "
        "8bit로 전송'(행 A-05 참조)."
    )
    note(doc, "'확정상태'가 [TBD]인 항목(범위·비트폭·주기 등)은 실제 센서/서보/RC 장비 데이터시트 및 기체 성능요구도(ORD) 확정 전까지의 예시값이다. '프로토콜 고정'은 MAVLink 표준 스케일/인코딩으로 이미 확정된 값이다. 배포 전 관련 부서(항전/성능/구조) 검토·확정이 반드시 필요하다.")

    start_landscape_section(doc)
    add_table(doc, PARAM_LIST_HEADERS, PARAM_LIST_ROWS, col_widths_cm=PARAM_LIST_COL_WIDTHS, repeat_header=True)
    end_landscape_section(doc)


def note(doc, text, color=RGBColor(0x99,0x00,0x00)):
    p = doc.add_paragraph()
    run = p.add_run("※ " + text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = color
